"""
AI-Powered Job Application Auto-Bidder
Uses Ollama to intelligently generate tailored resumes and answer application questions.
Now uses user_config.json and saves tailored resumes to organized folders.
"""

import requests
from bs4 import BeautifulSoup
import json
import ollama
from typing import Dict, List, Optional, Any
import re
import os
from datetime import datetime
from pathlib import Path
from docx import Document
from tailor_docx_resume import extract_resume_content, update_resume_sections
import copy


class AIJobBidder:
    """AI-powered job application system using Ollama."""
    
    def __init__(self, config_path="config.json"):
        """
        Initialize the AI Job Bidder.
        
        Args:
            config_path: Path to unified configuration JSON file
        """
        self.config = self.load_config(config_path)
        self.user_config = self.config['user_info']
        self.ai_settings = self.config['ai_settings']
        self.prompts = self.config['prompts']
        self.model = self.ai_settings['model']
        self.resumes_dir = Path("resumes")
        self.original_resume_dir = self.resumes_dir / "original"
        self.tailored_resume_dir = self.resumes_dir / "tailored"
        
        # Ensure directories exist
        self.tailored_resume_dir.mkdir(parents=True, exist_ok=True)
        
        print(f"Initializing AI Job Bidder")
        print(f"  Model: {self.model}")
        print(f"  User: {self.user_config['personal_info']['name']}")
        print(f"  Email: {self.user_config['personal_info']['email']}")
    
    def load_config(self, config_path: str) -> Dict:
        """Load user configuration from JSON file."""
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"❌ Config file not found: {config_path}")
            print("   Please create user_config.json with your information")
            raise
        except json.JSONDecodeError as e:
            print(f"❌ Invalid JSON in config file: {e}")
            raise
    
    def load_original_resume(self) -> tuple:
        """
        Load the original resume from file.
        
        Returns:
            Tuple of (resume_text, document_object or None)
            - If DOCX: (text_content, Document object)
            - If TXT: (text_content, None)
        """
        resume_path = self.user_config['files']['original_resume_path']
        try:
            if resume_path.endswith('.docx'):
                doc = Document(resume_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                return text, doc
            else:
                with open(resume_path, 'r', encoding='utf-8') as f:
                    return f.read(), None
        except FileNotFoundError:
            print(f"❌ Resume file not found: {resume_path}")
            raise
    
    def extract_docx_sections(self, doc: Document) -> Dict:
        """
        Extract sections from DOCX document for targeted tailoring.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Dict with section information including text and paragraph indices
        """
        sections = {
            'header': {'start': 0, 'end': 0, 'text': ''},
            'summary': {'start': -1, 'end': -1, 'text': ''},
            'education': {'start': -1, 'end': -1, 'text': ''},
            'skills': {'start': -1, 'end': -1, 'text': ''},
            'certifications': {'start': -1, 'end': -1, 'text': ''},
            'experience': {'start': -1, 'end': -1, 'jobs': []}
        }
        
        current_section = None
        section_headings = {
            'summary': ['summary of qualifications', 'summary', 'qualifications', 'profile'],
            'education': ['education', 'academic background'],
            'skills': ['technical skills', 'skills', 'core competencies'],
            'certifications': ['certifications', 'certificates'],
            'experience': ['work experience', 'professional experience', 'relevant work experience', 'experience']
        }
        
        # Track current job in experience section
        current_job = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            text_lower = text.lower()
            
            if not text:
                continue
            
            # Check if this is a section heading
            is_heading = False
            for section_name, keywords in section_headings.items():
                if any(text_lower == keyword or text_lower.startswith(keyword) for keyword in keywords):
                    current_section = section_name
                    sections[section_name]['start'] = i
                    is_heading = True
                    
                    # Close previous section
                    for sec in sections.values():
                        if isinstance(sec, dict) and 'end' in sec and sec['end'] == -1 and sec['start'] >= 0:
                            if sec['start'] < i:
                                sec['end'] = i - 1
                    break
            
            if is_heading:
                continue
            
            # Handle header (before any section)
            if current_section is None:
                sections['header']['text'] += text + '\n'
                sections['header']['end'] = i
                continue
            
            # Add content to current section
            if current_section == 'experience':
                # Check if this is a new job entry (company line or job title line)
                # Heuristic: if line has tab or looks like "Company    Location" or is a job title with dates
                if '\t' in text or ('–' in text and any(month in text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'])):
                    # This is either company/location or title/dates line
                    if current_job is None or 'dates' in current_job:
                        # Start new job
                        if current_job:
                            sections['experience']['jobs'].append(current_job)
                        current_job = {'company': '', 'title': '', 'dates': '', 'bullets': [], 'start_para': i}
                    
                    # Determine if it's company or title line
                    if '–' in text and any(month in text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                        # This is title and dates
                        parts = text.split('\t') if '\t' in text else [text]
                        current_job['title'] = parts[0].strip()
                        current_job['dates'] = parts[-1].strip()
                    else:
                        # This is company and location
                        parts = text.split('\t')
                        current_job['company'] = parts[0].strip()
                        current_job['location'] = parts[-1].strip() if len(parts) > 1 else ''
                elif current_job is not None:
                    # This is a bullet point
                    current_job['bullets'].append(text)
                    current_job['end_para'] = i
            else:
                sections[current_section]['text'] += text + '\n'
                sections[current_section]['end'] = i
        
        # Add last job if exists
        if current_job:
            sections['experience']['jobs'].append(current_job)
            sections['experience']['end'] = current_job.get('end_para', sections['experience']['start'])
        
        # Close any remaining open sections
        for sec in sections.values():
            if isinstance(sec, dict) and 'end' in sec and sec['end'] == -1 and sec['start'] >= 0:
                sec['end'] = len(doc.paragraphs) - 1
        
        return sections
    
    def save_tailored_resume(self, company: str, job_title: str, tailored_content, original_doc=None) -> str:
        """
        Save tailored resume to file in organized directory structure.
        Structure: tailored/{company}_{job_title}/{original_resume_name}.docx
        
        Args:
            company: Company name
            job_title: Job title
            tailored_content: Tailored resume (Document object or text string)
            original_doc: Original Document object (not used if tailored_content is Document)
            
        Returns:
            Path to saved resume file
        """
        # Create safe folder name from company and job title
        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')
        folder_name = f"{safe_company}_{safe_title}"
        
        # Create company/job folder inside tailored directory
        job_folder = self.tailored_resume_dir / folder_name
        job_folder.mkdir(parents=True, exist_ok=True)
        
        # Get original resume filename (without path)
        original_resume_path = Path(self.user_config['files']['original_resume_path'])
        original_filename = original_resume_path.name
        
        # Check if it's a Document object by checking for 'save' method
        if hasattr(tailored_content, 'save') and hasattr(tailored_content, 'paragraphs'):
            # Save as DOCX using original filename
            filepath = job_folder / original_filename
            tailored_content.save(str(filepath))
        else:
            # Text content - save using original filename with appropriate extension
            if original_doc:
                # Try to save as DOCX (legacy simple approach)
                filepath = job_folder / original_filename
                
                new_doc = copy.deepcopy(original_doc)
                resume_lines = tailored_content.split('\n')
                
                # Clear and rebuild
                while len(new_doc.paragraphs) > 0:
                    p = new_doc.paragraphs[0]
                    p._element.getparent().remove(p._element)
                
                for line in resume_lines:
                    if line.strip():
                        new_doc.add_paragraph(line)
                
                new_doc.save(str(filepath))
            else:
                # Save as TXT using original filename
                txt_filename = original_filename.replace('.docx', '.txt')
                filepath = job_folder / txt_filename
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(tailored_content)
        
        return str(filepath)
    
    def extract_application_form(self, job_url: str) -> Optional[Dict]:
        """
        Extract application form structure from job posting URL.
        
        Args:
            job_url: URL of the job posting
            
        Returns:
            Dictionary containing form fields and job details
        """
        try:
            print(f"\n{'='*70}")
            print(f"Extracting application form from: {job_url}")
            print(f"{'='*70}\n")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
            }
            
            response = requests.get(job_url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Extract from window.__appData
            script = soup.find('script', string=lambda t: t and 'window.__appData' in t)
            if not script:
                print("❌ Could not find application form data")
                return None
            
            # Parse the JSON (handle extra data at the end)
            content = script.string
            json_start = content.find('{')
            # Find the end of the JSON object by matching braces
            brace_count = 0
            json_end = json_start
            for i in range(json_start, len(content)):
                if content[i] == '{':
                    brace_count += 1
                elif content[i] == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        json_end = i + 1
                        break
            
            data = json.loads(content[json_start:json_end])
            
            # Extract job posting details
            posting = data.get('posting', {})
            form_data = posting.get('applicationForm', {})
            
            form_fields = []
            for entry in form_data.get('fieldEntries', []):
                field = entry.get('field', {})
                
                # Extract description
                desc_html = entry.get('descriptionHtml', '')
                description = ""
                if desc_html:
                    desc_soup = BeautifulSoup(desc_html, 'html.parser')
                    description = desc_soup.get_text(strip=True)
                
                # Extract options for select fields
                options = []
                if field.get('type') in ['ValueSelect', 'MultiValueSelect']:
                    select_options = field.get('selectableValues', [])
                    for opt in select_options:
                        if isinstance(opt, dict):
                            options.append(opt.get('label', opt.get('value', '')))
                        else:
                            options.append(str(opt))
                
                form_fields.append({
                    'title': field.get('title'),
                    'path': field.get('path'),
                    'type': field.get('type'),
                    'required': entry.get('isRequired', False),
                    'description': description,
                    'options': options if options else None
                })
            
            result = {
                'job_url': job_url,
                'company': data.get('organization', {}).get('name'),
                'job_title': posting.get('title'),
                'job_description': posting.get('descriptionPlainText', ''),
                'location': posting.get('locationName'),
                'form_fields': form_fields,
                'organization_id': posting.get('organizationId'),
                'posting_id': posting.get('id')
            }
            
            print(f"✅ Found {len(form_fields)} form fields")
            print(f"   Company: {result['company']}")
            print(f"   Position: {result['job_title']}")
            print(f"   Location: {result['location']}\n")
            
            return result
            
        except Exception as e:
            print(f"❌ Error extracting form: {e}")
            return None
    
    def tailor_resume(self, resume_text: str, job_description: str, job_title: str, 
                     company: str = "", doc: Document = None) -> str:
        """
        Use AI to tailor resume for specific job.
        
        Args:
            resume_text: Original resume content
            job_description: Full job description
            job_title: Job title
            company: Company name
            doc: Optional Document object for section-based tailoring
            
        Returns:
            Tailored resume text
        """
        print(f"\n{'='*70}")
        print("🤖 AI is tailoring your resume...")
        print(f"{'='*70}\n")
        
        # If we have a Document object, do section-by-section tailoring
        if doc:
            return self.tailor_docx_resume(doc, job_description, job_title, company)
        
        # Otherwise, do full text tailoring (legacy)
        prompt = f"""You are a professional resume writer. Rewrite the following resume to perfectly match this job posting.

JOB TITLE: {job_title}

JOB DESCRIPTION:
{job_description[:2000]}

ORIGINAL RESUME:
{resume_text}

Instructions:
1. Keep all factual information accurate (don't make up experience)
2. Emphasize skills and experiences that match the job requirements
3. Use keywords from the job description
4. Make it ATS-friendly
5. Keep it concise and professional
6. Format as plain text

Provide ONLY the tailored resume text, no additional commentary."""

        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            
            tailored_resume = response['message']['content'].strip()
            print(f"✅ Resume tailored successfully ({len(tailored_resume)} characters)\n")
            return tailored_resume
            
        except Exception as e:
            print(f"❌ Error tailoring resume: {e}")
            print("   Using original resume as fallback\n")
            return resume_text
    
    def tailor_docx_resume(self, doc: Document, job_description: str, 
                          job_title: str, company: str) -> Document:
        """
        Tailor DOCX resume by updating specific sections while preserving formatting.
        
        Args:
            doc: Original Document object
            job_description: Full job description
            job_title: Job title
            company: Company name
            
        Returns:
            Document object with tailored content
        """
        print("🔍 Extracting resume sections...")
        content = extract_resume_content(doc)
        
        # Tailor Summary section
        tailored_summary = ""
        if content['summary']['paras']:
            print("\n📝 Tailoring SUMMARY section...")
            original_summary = '\n'.join([p['text'] for p in content['summary']['paras']])
            tailored_summary = self.tailor_summary(
                original_summary,
                job_description,
                job_title,
                company
            )
            print("   ✅ Summary updated")
        
        # Tailor Skills section
        tailored_skills = ""
        if content['skills']['paras']:
            print("\n🛠️  Tailoring TECHNICAL SKILLS section...")
            original_skills = '\n'.join([p['text'] for p in content['skills']['paras']])
            tailored_skills = self.tailor_skills(
                original_skills,
                job_description,
                job_title,
                company
            )
            print("   ✅ Skills updated")
        
        # Tailor Experience section (each job)
        tailored_jobs = []
        if content['experience_jobs']:
            print(f"\n💼 Tailoring WORK EXPERIENCE ({len(content['experience_jobs'])} jobs)...")
            for idx, job in enumerate(content['experience_jobs'], 1):
                print(f"\n   Job {idx}/{len(content['experience_jobs'])}: {job.get('company', 'Unknown')}")
                original_bullets = [p['text'] for p in job['paras']]
                tailored_bullets = self.tailor_experience_job(
                    {
                        'company': job['company'],
                        'title': job['title'],
                        'dates': job['dates'],
                        'bullets': original_bullets
                    },
                    job_description,
                    job_title,
                    company
                )
                tailored_jobs.append({'bullets': tailored_bullets})
                print(f"      ✅ Updated {len(tailored_bullets)} bullet points")
        
        # Update the document
        print(f"\n📝 Applying changes to document...")
        doc = update_resume_sections(doc, content, tailored_summary, tailored_skills, tailored_jobs)
        
        # Validate and enhance resume
        print(f"\n🔍 Validating tailored resume...")
        doc = self.validate_and_enhance_resume(doc, content, tailored_summary, tailored_skills, 
                                                 tailored_jobs, job_description, job_title, company)
        
        print(f"\n✅ Resume tailoring complete!\n")
        return doc
    
    def tailor_summary(self, original_summary: str, job_description: str, 
                      job_title: str, company: str) -> str:
        """Use AI to tailor the summary section."""
        prompt_template = self.prompts['resume_tailoring']['summary_prompt']
        prompt = prompt_template.format(
            job_title=job_title,
            company=company,
            job_description=job_description[:1500],
            original_summary=original_summary
        )
        
        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            return response['message']['content'].strip()
        except Exception as e:
            print(f"      ⚠️  Error: {e}, using original")
            return original_summary
    
    def tailor_skills(self, original_skills: str, job_description: str,
                     job_title: str, company: str) -> str:
        """Use AI to tailor the skills section."""
        prompt_template = self.prompts['resume_tailoring']['skills_prompt']
        prompt = prompt_template.format(
            job_title=job_title,
            company=company,
            job_description=job_description[:1500],
            original_skills=original_skills
        )
        
        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            return response['message']['content'].strip()
        except Exception as e:
            print(f"      ⚠️  Error: {e}, using original")
            return original_skills
    
    def tailor_experience_job(self, job: Dict, job_description: str,
                             target_job_title: str, target_company: str) -> List[str]:
        """Use AI to tailor bullet points for a specific job entry."""
        prompt_template = self.prompts['resume_tailoring']['experience_prompt']
        
        bullets_text = '\n'.join(job['bullets'])
        num_bullets = len(job['bullets'])
        
        prompt = prompt_template.format(
            job_title=target_job_title,
            company=target_company,
            job_description=job_description[:1500],
            exp_company=job.get('company', ''),
            exp_title=job.get('title', ''),
            exp_dates=job.get('dates', ''),
            original_bullets=bullets_text,
            num_bullets=num_bullets
        )
        
        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            # Split response into individual bullets
            tailored_text = response['message']['content'].strip()
            bullets = [b.strip() for b in tailored_text.split('\n') if b.strip()]
            
            # Ensure we return the same number of bullets
            if len(bullets) != num_bullets:
                print(f"      ⚠️  Warning: Got {len(bullets)} bullets, expected {num_bullets}")
                # Pad or truncate to match
                if len(bullets) < num_bullets:
                    bullets.extend(job['bullets'][len(bullets):])  # Add original bullets
                else:
                    bullets = bullets[:num_bullets]
            
            return bullets
        except Exception as e:
            print(f"      ⚠️  Error: {e}, using original")
            return job['bullets']
    
    def validate_and_enhance_resume(self, doc: Document, content: Dict, 
                                     tailored_summary: str, tailored_skills: str,
                                     tailored_jobs: List[Dict], job_description: str,
                                     job_title: str, company: str) -> Document:
        """
        Validate tailored resume and enhance if sections are missing or incomplete.
        
        Args:
            doc: Document object
            content: Extracted resume content structure
            tailored_summary: Tailored summary text
            tailored_skills: Tailored skills text  
            tailored_jobs: List of tailored job entries
            job_description: Target job description
            job_title: Target job title
            company: Target company
            
        Returns:
            Enhanced Document object
        """
        issues_found = []
        
        # Check Summary section
        if not tailored_summary or len(tailored_summary.strip()) < 50:
            issues_found.append("Summary too short or missing")
            print("   ⚠️  Summary section incomplete")
        
        # Check Skills section
        if not tailored_skills or len(tailored_skills.strip()) < 20:
            issues_found.append("Skills section incomplete")
            print("   ⚠️  Skills section incomplete")
        else:
            # Check if all 4 categories are present
            required_categories = ['Programming & Frameworks:', 'Data & AI/ML Tools:', 
                                   'Cloud & DevOps:', 'Databases & APIs:']
            missing_categories = [cat for cat in required_categories if cat not in tailored_skills]
            if missing_categories:
                issues_found.append(f"Missing skill categories: {missing_categories}")
                print(f"   ⚠️  Missing categories: {', '.join(missing_categories)}")
        
        # Check Experience section
        if not tailored_jobs or len(tailored_jobs) == 0:
            issues_found.append("Experience section missing")
            print("   ⚠️  Experience section incomplete")
        else:
            for idx, job in enumerate(tailored_jobs, 1):
                if not job.get('bullets') or len(job.get('bullets', [])) == 0:
                    issues_found.append(f"Job {idx} has no bullets")
                    print(f"   ⚠️  Job {idx} missing bullet points")
        
        # If issues found, attempt to re-enhance
        if issues_found:
            print(f"\n🔧 Enhancing incomplete sections...")
            
            # Re-enhance summary if needed
            if any('Summary' in issue for issue in issues_found):
                print("   🔄 Re-generating summary...")
                original_summary = '\n'.join([p['text'] for p in content['summary']['paras']])
                tailored_summary = self.tailor_summary(original_summary, job_description, job_title, company)
                # Re-apply to document
                content_updated = extract_resume_content(doc)
                doc = update_resume_sections(doc, content_updated, tailored_summary, tailored_skills, tailored_jobs)
            
            # Re-enhance skills if needed
            if any('Skills' in issue or 'categories' in issue for issue in issues_found):
                print("   🔄 Re-generating skills...")
                original_skills = '\n'.join([p['text'] for p in content['skills']['paras']])
                tailored_skills = self.tailor_skills(original_skills, job_description, job_title, company)
                # Re-apply to document
                content_updated = extract_resume_content(doc)
                doc = update_resume_sections(doc, content_updated, tailored_summary, tailored_skills, tailored_jobs)
            
            print("   ✅ Enhancement complete")
        else:
            print("   ✅ All sections validated successfully")
        
        return doc
    
    def replace_section_text(self, doc: Document, section: Dict, new_text: str):
        """Replace text in a document section while preserving formatting."""
        start = section['start']
        end = section['end']
        
        if start < 0 or end < 0:
            return
        
        # Skip the heading, start from content
        content_start = start + 1
        
        # Clear existing content paragraphs
        paragraphs_to_remove = []
        for i in range(content_start, end + 1):
            if i < len(doc.paragraphs):
                paragraphs_to_remove.append(doc.paragraphs[i])
        
        for para in paragraphs_to_remove:
            p_element = para._element
            p_element.getparent().remove(p_element)
        
        # Insert new content after the heading
        new_lines = [line.strip() for line in new_text.split('\n') if line.strip()]
        heading_para = doc.paragraphs[start]
        
        for line in reversed(new_lines):
            # Insert after heading
            new_para = heading_para._element.getparent().insert(
                heading_para._element.getparent().index(heading_para._element) + 1,
                copy.deepcopy(doc.paragraphs[content_start]._element if content_start < len(doc.paragraphs) else heading_para._element)
            )
            # Update text
            new_para_obj = None
            for p in doc.paragraphs:
                if p._element == new_para:
                    new_para_obj = p
                    break
            if new_para_obj:
                new_para_obj.text = line
    
    def replace_job_bullets(self, doc: Document, job: Dict, new_bullets: List[str]):
        """Replace bullet points for a job entry while preserving formatting."""
        if 'start_para' not in job or 'end_para' not in job:
            return
        
        start = job['start_para']
        end = job['end_para']
        
        # Find where bullets start (skip company and title lines)
        bullets_start = start
        for i in range(start, end + 1):
            if i < len(doc.paragraphs):
                text = doc.paragraphs[i].text.strip()
                # Skip company/location and title/dates lines
                if '\t' in text or ('–' in text and any(m in text for m in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'])):
                    bullets_start = i + 1
                else:
                    break
        
        # Remove old bullet paragraphs
        paragraphs_to_remove = []
        for i in range(bullets_start, end + 1):
            if i < len(doc.paragraphs):
                paragraphs_to_remove.append(doc.paragraphs[i])
        
        for para in paragraphs_to_remove:
            p_element = para._element
            p_element.getparent().remove(p_element)
        
        # Insert new bullets
        if bullets_start > 0 and bullets_start - 1 < len(doc.paragraphs):
            ref_para = doc.paragraphs[bullets_start - 1]
            for bullet in reversed(new_bullets):
                new_para = ref_para._element.getparent().insert(
                    ref_para._element.getparent().index(ref_para._element) + 1,
                    copy.deepcopy(ref_para._element)
                )
                # Update text
                for p in doc.paragraphs:
                    if p._element == new_para:
                        p.text = bullet
                        break
    
    def answer_question(self, question: str, job_description: str, 
                       job_title: str, company: str, user_background: str) -> str:
        """
        Use AI to intelligently answer application questions.
        Uses specific prompts based on question type.
        
        Args:
            question: The question to answer
            job_description: Full job description
            job_title: Job title
            company: Company name
            user_background: Brief user background/resume
            
        Returns:
            AI-generated answer
        """
        print(f"   🤖 Generating answer for: {question[:60]}...")
        
        # Detect question type and use appropriate prompt
        question_lower = question.lower()
        
        if 'how did you find' in question_lower or 'where did you hear' in question_lower:
            prompt_template = self.prompts['application_questions'].get('how_did_you_find_us_prompt')
            prompt = prompt_template.format(company=company, job_title=job_title)
        elif 'why' in question_lower and ('interest' in question_lower or 'want to work' in question_lower or 'attracted' in question_lower):
            prompt_template = self.prompts['application_questions'].get('why_interested_prompt')
            prompt = prompt_template.format(
                company=company,
                job_title=job_title,
                job_description=job_description[:1500],
                user_background=user_background[:500]
            )
        elif 'achievement' in question_lower or 'accomplishment' in question_lower or 'proud of' in question_lower:
            prompt_template = self.prompts['application_questions'].get('recent_achievement_prompt')
            prompt = prompt_template.format(
                company=company,
                job_title=job_title,
                job_description=job_description[:1500],
                user_background=user_background[:500]
            )
        else:
            # Default prompt
            prompt_template = self.prompts['application_questions']['default_answer_prompt']
            prompt = prompt_template.format(
                company=company,
                job_title=job_title,
                job_description=job_description[:1500],
                user_background=user_background[:500],
                question=question
            )

        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            
            answer = response['message']['content'].strip()
            print(f"      ✅ Generated ({len(answer)} characters)")
            return answer
            
        except Exception as e:
            print(f"      ❌ Error generating answer: {e}")
            return f"I am very interested in this {job_title} position at {company}."
    
    def select_best_option(self, question: str, options: List[str], 
                          job_description: str, job_title: str, company: str,
                          user_background: str, multi_select: bool = False) -> str:
        """
        Use AI to select the best option(s) from a list based on context.
        
        Args:
            question: The question/field title
            options: List of available options
            job_description: Full job description
            job_title: Job title
            company: Company name
            user_background: User background
            multi_select: If True, can select multiple options
            
        Returns:
            Selected option(s) as string or list
        """
        print(f"   🤖 Selecting from {len(options)} options for: {question[:60]}...")
        
        options_text = '\n'.join([f"- {opt}" for opt in options])
        
        if multi_select:
            instruction = "Select ALL applicable options from the list (can be multiple). Return as comma-separated values."
        else:
            instruction = "Select the SINGLE BEST option from the list that matches. Return ONLY the exact option text."
        
        prompt = f"""You are helping fill out a job application form.

COMPANY: {company}
JOB TITLE: {job_title}

QUESTION/FIELD: {question}

AVAILABLE OPTIONS:
{options_text}

APPLICANT BACKGROUND:
{user_background[:800]}

JOB REQUIREMENTS:
{job_description[:800]}

INSTRUCTIONS:
{instruction}

Based on the applicant's background and job requirements, which option(s) best match?

Output format: Return ONLY the option text exactly as shown, nothing else."""
        
        try:
            response = ollama.chat(model=self.model, messages=[
                {'role': 'user', 'content': prompt}
            ])
            
            answer = response['message']['content'].strip()
            
            # Parse response
            if multi_select:
                # Split by commas and clean up
                selected = [opt.strip() for opt in answer.split(',')]
                # Validate options exist
                valid_selected = [s for s in selected if any(s.lower() in opt.lower() or opt.lower() in s.lower() for opt in options)]
                return valid_selected if valid_selected else [options[0]]
            else:
                # Find best matching option
                for opt in options:
                    if opt.lower() in answer.lower() or answer.lower() in opt.lower():
                        return opt
                # Fallback to first option
                return options[0]
                
        except Exception as e:
            print(f"      ❌ Error selecting option: {e}")
            return options[0] if not multi_select else [options[0]]
    
    def fill_application(self, form_data: Dict) -> Dict:
        """
        Fill application form with AI-generated content using user config.
        
        Args:
            form_data: Form structure from extract_application_form()
            
        Returns:
            Dictionary with filled form data
        """
        print(f"\n{'='*70}")
        print("📝 Filling application form with AI assistance...")
        print(f"{'='*70}\n")
        
        job_description = form_data['job_description']
        job_title = form_data['job_title']
        company = form_data['company']
        
        # Load original resume (returns tuple: text and doc object if DOCX)
        original_resume_text, original_doc = self.load_original_resume()
        
        # Get user config
        personal_info = self.user_config['personal_info']
        links = self.user_config['links']
        work_auth = self.user_config['work_authorization']
        demographics = self.user_config.get('demographics', {})
        background = self.user_config['background']
        
        filled_data = {
            'company': company,
            'job_title': job_title,
            'job_url': form_data['job_url'],
            'timestamp': datetime.now().isoformat(),
            'fields': {}
        }
        
        for field in form_data['form_fields']:
            field_title = field['title']
            field_type = field['type']
            field_path = field['path']
            required = field['required']
            
            print(f"Field: {field_title} ({field_type}){' *REQUIRED*' if required else ''}")
            
            # Handle different field types - CHECK SPECIFIC FIELDS FIRST
            if 'middle' in field_title.lower() and 'name' in field_title.lower():
                filled_data['fields'][field_path] = personal_info.get('middle_name', '')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
            
            elif 'pronoun' in field_title.lower():
                filled_data['fields'][field_path] = personal_info.get('pronouns', 'he/him/his')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
            
            elif field_path == '_systemfield_name' or ('name' in field_title.lower() and 'middle' not in field_title.lower()):
                filled_data['fields'][field_path] = personal_info['name']
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif field_path == '_systemfield_email' or 'email' in field_title.lower():
                filled_data['fields'][field_path] = personal_info['email']
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif 'phone' in field_title.lower():
                filled_data['fields'][field_path] = personal_info.get('phone', '')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif field_path == '_systemfield_resume' or 'resume' in field_title.lower():
                # COMMENTED OUT FOR TESTING - Use original resume directly
                # Tailor the resume (pass doc object for DOCX tailoring)
                # if original_doc:
                #     # DOCX: returns Document object
                #     tailored_doc = self.tailor_resume(
                #         original_resume_text, 
                #         job_description, 
                #         job_title,
                #         company,
                #         original_doc
                #     )
                #     # Extract text for the form field
                #     tailored_text = '\n'.join([p.text for p in tailored_doc.paragraphs if p.text.strip()])
                #     filled_data['fields'][field_path] = tailored_text
                #     
                #     # Save tailored DOCX
                #     resume_path = self.save_tailored_resume(company, job_title, tailored_doc)
                # else:
                #     # TXT: returns text string
                #     tailored_text = self.tailor_resume(
                #         original_resume_text,
                #         job_description,
                #         job_title,
                #         company
                #     )
                #     filled_data['fields'][field_path] = tailored_text
                #     resume_path = self.save_tailored_resume(company, job_title, tailored_text)
                
                # Use original resume path directly for testing
                resume_path = self.user_config['files']['original_resume_path']
                filled_data['fields'][field_path] = original_resume_text
                filled_data['tailored_resume_path'] = resume_path
                print(f"   ✅ Using original resume: {resume_path}\n")
                
            elif field_path == '_systemfield_location' or 'location' in field_title.lower():
                filled_data['fields'][field_path] = personal_info['location']
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif 'linkedin' in field_title.lower():
                filled_data['fields'][field_path] = links.get('linkedin', '')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif 'github' in field_title.lower():
                filled_data['fields'][field_path] = links.get('github', '')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif 'website' in field_title.lower() or 'portfolio' in field_title.lower():
                filled_data['fields'][field_path] = links.get('website', '')
                print(f"   ✅ Set to: {filled_data['fields'][field_path]}\n")
                
            elif field_type == 'Boolean':
                # For yes/no questions, use config values
                if 'authorized' in field_title.lower() and 'united states' in field_title.lower():
                    filled_data['fields'][field_path] = work_auth['authorized_to_work_us']
                    print(f"   ✅ Set to: {'Yes' if work_auth['authorized_to_work_us'] else 'No'}\n")
                elif 'authorized' in field_title.lower() and 'canada' in field_title.lower():
                    filled_data['fields'][field_path] = work_auth['authorized_to_work_canada']
                    print(f"   ✅ Set to: {'Yes' if work_auth['authorized_to_work_canada'] else 'No'}\n")
                elif 'visa' in field_title.lower() or 'sponsorship' in field_title.lower():
                    filled_data['fields'][field_path] = work_auth['needs_visa_sponsorship']
                    print(f"   ✅ Set to: {'Yes' if work_auth['needs_visa_sponsorship'] else 'No'}\n")
                else:
                    filled_data['fields'][field_path] = True
                    print(f"   ✅ Set to: Yes\n")
                    
            elif field_type in ['LongText', 'String']:
                # Use AI to answer custom questions
                if field['description']:
                    question = f"{field_title}: {field['description']}"
                else:
                    question = field_title
                
                answer = self.answer_question(
                    question, 
                    job_description, 
                    job_title, 
                    company, 
                    background['elevator_pitch']
                )
                filled_data['fields'][field_path] = answer
                print()
            
            elif field_type in ['ValueSelect', 'MultiValueSelect']:
                # Handle specific demographics and location fields first
                options = field.get('options', [])
                
                # Check for specific field types
                if 'pronoun' in field_title.lower():
                    # Map our pronoun format to their options
                    pronoun_input = personal_info.get('pronouns', 'he/him/his')
                    
                    # Map common formats to option text
                    pronoun_mapping = {
                        'he/him/his': 'He/him/his',
                        'she/her/hers': 'She/her/hers',
                        'they/them/theirs': 'They/them/theirs'
                    }
                    
                    # Try to find matching option
                    pronoun_value = pronoun_mapping.get(pronoun_input.lower(), pronoun_input)
                    
                    # Check if it's in the options
                    if options:
                        # Find closest match
                        for opt in options:
                            if pronoun_input.lower() in opt.lower() or opt.lower() in pronoun_input.lower():
                                pronoun_value = opt
                                break
                    
                    filled_data['fields'][field_path] = pronoun_value
                    print(f"   ✅ Selected: {pronoun_value}\n")
                elif 'gender' in field_title.lower():
                    # Use demographics gender
                    gender_value = demographics.get('gender', 'Man')

                    # Validate gender value against available options
                    if options:
                        print(f"   🔍 Available gender options: {options}")
                        if gender_value not in options:
                            print(f"   ⚠️  Gender value '{gender_value}' not in options. Defaulting to first available option.")
                            gender_value = options[0]  # Default to the first option if mismatch
                    else:
                        print("   ⚠️  No gender options available in the form.")

                    filled_data['fields'][field_path] = gender_value
                    print(f"   ✅ Selected: {gender_value}\n")
                elif 'race' in field_title.lower() or 'ethnicity' in field_title.lower():
                    # Use demographics race
                    race_value = demographics.get('race', 'White')
                    filled_data['fields'][field_path] = race_value
                    print(f"   ✅ Selected: {race_value}\n")
                elif 'disability' in field_title.lower() or 'disabled' in field_title.lower():
                    # Use demographics disability
                    disability_value = demographics.get('disability', 'No')
                    filled_data['fields'][field_path] = disability_value
                    print(f"   ✅ Selected: {disability_value}\n")
                elif 'veteran' in field_title.lower():
                    # Use demographics veteran_status
                    veteran_value = demographics.get('veteran_status', 'No')
                    filled_data['fields'][field_path] = veteran_value
                    print(f"   ✅ Selected: {veteran_value}\n")
                elif 'state' in field_title.lower() or 'province' in field_title.lower():
                    # Use state from personal info
                    state_value = personal_info.get('state', 'Massachusetts')
                    filled_data['fields'][field_path] = state_value
                    print(f"   ✅ Selected: {state_value}\n")
                elif 'have you ever been' in field_title.lower() or 'have you worked' in field_title.lower() or 'previously employed' in field_title.lower():
                    # For "Have you ever been" type questions, default to No
                    filled_data['fields'][field_path] = 'No'
                    print(f"   ✅ Selected: No (default for 'have you ever' questions)\n")
                elif options:
                    # Use AI to select best option(s) from available choices
                    if field['description']:
                        question = f"{field_title}: {field['description']}"
                    else:
                        question = field_title
                    
                    selected = self.select_best_option(
                        question,
                        options,
                        job_description,
                        job_title,
                        company,
                        background['elevator_pitch'],
                        field_type == 'MultiValueSelect'
                    )
                    filled_data['fields'][field_path] = selected
                    if isinstance(selected, list):
                        print(f"   ✅ Selected: {', '.join(selected)}\n")
                    else:
                        print(f"   ✅ Selected: {selected}\n")
                else:
                    print(f"   ⚠️  No options available for {field_type}, skipping\n")
            
            else:
                print(f"   ⚠️  Unknown field type, skipping\n")
        
        print(f"{'='*70}")
        print(f"✅ Application form filled successfully!")
        print(f"{'='*70}\n")
        
        return filled_data
    
    def generate_application_preview(self, filled_data: Dict) -> str:
        """
        Generate a human-readable preview of the filled application.
        
        Args:
            filled_data: Filled form data
            
        Returns:
            Formatted preview string
        """
        preview = f"""
{'='*70}
APPLICATION PREVIEW
{'='*70}

Company: {filled_data['company']}
Position: {filled_data['job_title']}
Job URL: {filled_data['job_url']}
Generated: {filled_data.get('timestamp', 'N/A')}

{'='*70}
FILLED FIELDS:
{'='*70}

"""
        for field_path, value in filled_data['fields'].items():
            # Clean up field names
            field_name = field_path.replace('_systemfield_', '').replace('_', ' ').title()
            
            if isinstance(value, bool):
                value_str = 'Yes' if value else 'No'
            elif len(str(value)) > 200:
                value_str = str(value)[:200] + '...'
            else:
                value_str = str(value)
            
            preview += f"{field_name}:\n{value_str}\n\n"
        
        if 'tailored_resume_path' in filled_data:
            preview += f"{'='*70}\n"
            preview += f"Tailored Resume: {filled_data['tailored_resume_path']}\n"
        
        preview += f"{'='*70}\n"
        
        return preview
    
    def analyze_page_with_vision(self, screenshot_path: str, fields_to_fill: Dict[str, Any]) -> Dict[str, str]:
        """Use vision model to analyze page and get instructions for filling form."""
        print(f"   🔍 Using AI vision to analyze the page...\n")
        
        import base64
        
        # Convert screenshot to base64
        with open(screenshot_path, 'rb') as img_file:
            img_base64 = base64.b64encode(img_file.read()).decode('utf-8')
        
        prompt = """Analyze this job application form screenshot. List all visible input fields, text areas, file upload buttons, checkboxes, and dropdown menus you can see.

For each field, provide:
- Field label/text visible on the form
- Element type (text input, textarea, file upload, checkbox, dropdown)
- Any unique identifiers visible (placeholder text, nearby text)

Be concise. Focus on form elements only."""

        try:
            response = requests.post(
                'http://localhost:11434/api/generate',
                json={
                    'model': 'llama3.2-vision',
                    'prompt': prompt,
                    'images': [img_base64],
                    'stream': False,
                    'options': {
                        'temperature': 0.1,
                        'num_predict': 500
                    }
                },
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()['response']
                print(f"   ✅ Vision analysis:\n{result[:500]}...\n")
                return {'analysis': result}
            else:
                print(f"   ❌ Vision analysis failed\n")
                return {}
        except Exception as e:
            print(f"   ⚠️  Vision analysis error: {str(e)[:100]}\n")
            return {}
    
    def submit_application_with_browser(self, filled_data: Dict, form_data: Dict, 
                                        headless: bool = False) -> Dict:
        """
        Submit filled application using Playwright browser automation.
        
        Args:
            filled_data: Filled application data
            form_data: Original form data with job URL
            headless: If True, run browser in headless mode
            
        Returns:
            Dict with submission status
        """
        print(f"\n{'='*70}")
        print("🤖 Submitting application with browser automation...")
        print(f"{'='*70}\n")
        
        try:
            from playwright.sync_api import sync_playwright
            import time
            
            job_url = form_data['job_url']
            resume_path = filled_data.get('tailored_resume_path', '')
            
            print(f"   🌐 Opening browser...")
            print(f"   📄 Resume: {resume_path}\n")
            
            with sync_playwright() as p:
                # Launch real Chrome browser with maximum stealth
                print(f"   🌐 Opening Chrome browser...")
                
                # Use persistent context with a profile to look like a real user
                import tempfile
                import shutil
                
                # OPTION: Use your real Chrome profile instead of temp (more realistic)
                # Uncomment this to use real profile (WARNING: may affect your real browser)
                # profile_dir = os.path.expanduser("~/Library/Application Support/Google/Chrome/Default")
                
                # Create a temporary profile directory
                profile_dir = tempfile.mkdtemp(prefix="chrome_profile_")
                print(f"   📁 Using Chrome profile: {profile_dir}")
                
                # Anti-spam settings - fully automated
                WARM_UP_MODE = False  # Disabled - no manual browsing needed
                AUTO_SUBMIT = True  # Set to True to auto-submit after review
                
                try:
                    browser_context = p.chromium.launch_persistent_context(
                        user_data_dir=profile_dir,
                        channel="chrome",  # Real Chrome, not Chromium
                        headless=headless,
                        args=[
                            '--disable-blink-features=AutomationControlled',
                            '--disable-dev-shm-usage',
                            '--disable-infobars',
                            '--no-first-run',
                            '--no-service-autorun',
                            '--password-store=basic',
                        ],
                        viewport={'width': 1920, 'height': 1080},
                        user_agent='Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                        locale='en-US',
                        timezone_id='America/New_York',
                    )
                    
                    page = browser_context.pages[0] if browser_context.pages else browser_context.new_page()
                    
                    # Inject comprehensive anti-detection scripts BEFORE navigating
                    page.add_init_script("""
                        // Remove webdriver property
                        Object.defineProperty(navigator, 'webdriver', {
                            get: () => undefined
                        });
                        
                        // Override the automation flag
                        delete navigator.__proto__.webdriver;
                        
                        // Mock chrome object with more properties
                        window.chrome = {
                            runtime: {},
                            loadTimes: function() {},
                            csi: function() {},
                            app: {}
                        };
                        
                        // Mock permissions
                        const originalQuery = window.navigator.permissions.query;
                        window.navigator.permissions.query = (parameters) => (
                            parameters.name === 'notifications' ?
                                Promise.resolve({ state: Notification.permission }) :
                                originalQuery(parameters)
                        );
                        
                        // Mock plugins with realistic values
                        Object.defineProperty(navigator, 'plugins', {
                            get: () => [
                                {name: 'Chrome PDF Plugin', filename: 'internal-pdf-viewer', description: 'Portable Document Format'},
                                {name: 'Chrome PDF Viewer', filename: 'mhjfbmdgcfjbbpaeojofohoefgiehjai', description: ''},
                                {name: 'Native Client', filename: 'internal-nacl-plugin', description: ''}
                            ]
                        });
                        
                        // Mock languages
                        Object.defineProperty(navigator, 'languages', {
                            get: () => ['en-US', 'en']
                        });
                        
                        // Mock hardware concurrency
                        Object.defineProperty(navigator, 'hardwareConcurrency', {
                            get: () => 8
                        });
                        
                        // Mock device memory
                        Object.defineProperty(navigator, 'deviceMemory', {
                            get: () => 8
                        });
                        
                        // Override toString to hide proxy
                        const originalToString = Function.prototype.toString;
                        Function.prototype.toString = function() {
                            if (this === navigator.permissions.query) {
                                return 'function query() { [native code] }';
                            }
                            return originalToString.apply(this, arguments);
                        };
                        
                        // Mock WebGL vendor/renderer
                        const getParameter = WebGLRenderingContext.prototype.getParameter;
                        WebGLRenderingContext.prototype.getParameter = function(parameter) {
                            if (parameter === 37445) {
                                return 'Intel Inc.';
                            }
                            if (parameter === 37446) {
                                return 'Intel Iris OpenGL Engine';
                            }
                            return getParameter.apply(this, [parameter]);
                        };
                        
                        // Add realistic screen properties
                        Object.defineProperty(screen, 'availTop', {
                            get: () => 0
                        });
                        
                        // Mock connection
                        Object.defineProperty(navigator, 'connection', {
                            get: () => ({
                                effectiveType: '4g',
                                rtt: 50,
                                downlink: 10,
                                saveData: false
                            })
                        });
                    """)
                    
                except Exception as e:
                    print(f"   ⚠️  Could not use persistent context: {e}")
                    # Cleanup temp dir if failed
                    if profile_dir:
                        shutil.rmtree(profile_dir, ignore_errors=True)
                    raise
                
                # Navigate to job posting with human-like behavior
                print(f"   ⏳ Loading job page...")
                
                import random
                
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        page.goto(job_url, timeout=60000)
                        page.wait_for_load_state('networkidle', timeout=30000)
                        
                        break
                    except Exception as e:
                        if attempt < max_retries - 1:
                            print(f"   ⚠️  Attempt {attempt + 1} failed, retrying...")
                        else:
                            raise
                
                # Look for and click "Apply" button if it exists
                print(f"   🔍 Looking for Apply button...")
                
                apply_button_selectors = [
                    'button:has-text("Apply")',
                    'a:has-text("Apply")',
                    '[data-testid="apply-button"]',
                    '.apply-button',
                    'button:has-text("Submit Application")'
                ]
                
                apply_clicked = False
                for selector in apply_button_selectors:
                    try:
                        if page.locator(selector).count() > 0:
                            button = page.locator(selector).first
                            button.click()
                            print(f"   ✅ Clicked Apply button\n")
                            apply_clicked = True
                            page.wait_for_load_state('networkidle')
                            break
                    except:
                        continue
                
                if not apply_clicked:
                    print(f"   ℹ️  No Apply button found (form may already be visible)\n")
                
                # Fill form fields
                print(f"   📝 Filling form fields...\n")
                
                # Wait for form to be fully loaded
                form_fields_found = False
                try:
                    page.wait_for_selector('input, textarea, select', timeout=5000)
                    form_fields_found = True
                except:
                    pass
                
                # If no form fields, look for Apply/Submit button to reveal the form
                if not form_fields_found and not apply_clicked:
                    print(f"   ⚠️  No form fields detected, looking for Apply button...\n")
                    
                    apply_selectors = [
                        'button:has-text("Apply for this job")',
                        'button:has-text("Apply")',
                        'a:has-text("Apply for this job")',
                        'a:has-text("Apply")',
                        '[data-testid*="apply"]',
                        'button[type="button"]:has-text("Apply")',
                        '.apply-button',
                        '#apply-button'
                    ]
                    
                    for selector in apply_selectors:
                        try:
                            if page.locator(selector).count() > 0:
                                print(f"   🎯 Found Apply button with selector: {selector}")
                                page.locator(selector).first.click()
                                print(f"   ✅ Clicked Apply button, waiting for form...\n")
                                page.wait_for_load_state('networkidle')
                                
                                # Check again for form fields
                                try:
                                    page.wait_for_selector('input, textarea, select', timeout=5000)
                                    form_fields_found = True
                                    print(f"   ✅ Form fields now visible!\n")
                                except:
                                    pass
                                break
                        except Exception as e:
                            continue
                
                if not form_fields_found:
                    print(f"   ⚠️  Still no form fields detected after trying Apply button\n")
                
                # Take screenshot for debugging
                debug_screenshot = f"form_debug_{int(time.time())}.png"
                page.screenshot(path=debug_screenshot, full_page=True)
                print(f"   📸 Debug screenshot saved: {debug_screenshot}\n")
                
                # Track which boolean buttons we've clicked (by button text)
                boolean_yes_clicks = 0
                boolean_no_clicks = 0
                
                for field in form_data['form_fields']:
                    field_path = field['path']
                    field_title = field['title']
                    field_type = field['type']
                    value = filled_data['fields'].get(field_path)
                    
                    print(f"\n      📋 Processing: {field_title} (Type: {field_type})")
                    
                    if value is None:
                        print(f"         ⚠️  No value provided, skipping")
                        continue
                    
                    print(f"         📝 Value: {str(value)[:100]}...")
                    
                    try:
                        element = None
                        
                        # Try different selector strategies based on field type
                        if field_type == 'File':
                            selectors = [
                                'input[type="file"]',
                                '[data-testid*="resume"]',
                                '[data-testid*="upload"]',
                                'button:has-text("Upload") + input[type="file"]'
                            ]
                        elif field_type == 'Boolean':
                            # For checkboxes, use name attribute with lowercase field_path
                            selectors = [
                                f'input[type="checkbox"][name="{field_path.lower()}"]',
                                f'input[type="checkbox"][name="{field_path}"]',
                                f'input[type="checkbox"][id*="{field_path.lower()}"]',
                            ]
                        elif field_type == 'Location':
                            # Location fields use special autocomplete - look for placeholder "Start typing"
                            print(f"         🔍 Looking for Location field...")
                            selectors = [
                                'input[placeholder="Start typing..."]',
                                'input[placeholder*="Start typing" i]',
                                'input[placeholder*="location" i]',
                                'input[placeholder*="city" i]',
                                'input[aria-label*="location" i]',
                                f'[id="{field_path.lower()}"]' if field_path else None,
                                f'[name="{field_path.lower()}"]' if field_path else None,
                            ]
                            selectors = [s for s in selectors if s]  # Remove None values
                        else:
                            # Try multiple selector patterns
                            selectors = [
                                f'input[placeholder*="{field_title}" i]',
                                f'textarea[placeholder*="{field_title}" i]',
                                f'input[aria-label*="{field_title}" i]',
                                f'textarea[aria-label*="{field_title}" i]',
                            ]
                            
                            # Add ID-based selectors with lowercase
                            if field_path:
                                selectors.extend([
                                    f'[id="{field_path.lower()}"]',
                                    f'[name="{field_path.lower()}"]',
                                    f'[id="{field_path}"]',
                                    f'[name="{field_path}"]',
                                ])
                        
                        # Try each selector
                        for selector in selectors:
                            try:
                                count = page.locator(selector).count()
                                if count > 0:
                                    print(f"         ✓ Found with selector: {selector} (count: {count})")
                                    element = page.locator(selector).first
                                    break
                            except:
                                continue
                        
                        # If still not found, try label-based approach
                        if not element:
                            try:
                                # Find by label text (case insensitive)
                                labels = page.locator(f'text="{field_title}" >> xpath=..');
                                if labels.count() > 0:
                                    # Try to find nearby input
                                    parent = labels.first
                                    nearby_inputs = parent.locator('input, textarea, select')
                                    if nearby_inputs.count() > 0:
                                        element = nearby_inputs.first
                            except:
                                pass
                        
                        if element and element.count() > 0:
                            if field_type == 'File':
                                # Handle file upload
                                if resume_path and Path(resume_path).exists():
                                    element.set_input_files(str(Path(resume_path).absolute()))
                                    print(f"      ✅ Uploaded: {field_title}")
                            elif field_type == 'Boolean':
                                # Handle checkbox/boolean - these are Yes/No button pairs
                                button_text = "Yes" if value else "No"
                                
                                try:
                                    # First: Set checkbox value via JS (hidden input)
                                    element.evaluate(f'el => el.checked = {str(value).lower()}')
                                    print(f"      ✅ Set checkbox via JS: {field_title}")
                                    
                                    # Second: Find the button that's NEAR this checkbox input
                                    # Strategy: Find parent container, then look for button within it
                                    print(f"         🔘 Looking for '{button_text}' button near this field...")
                                    
                                    try:
                                        # Get the parent container of the checkbox
                                        parent = element.locator('xpath=ancestor::div[contains(@class, "css-") or contains(@class, "field")]').first
                                        
                                        # Look for the button within this parent
                                        button = parent.locator(f'button:has-text("{button_text}")').first
                                        
                                        if button.count() > 0:
                                            button.scroll_into_view_if_needed()
                                            button.click(timeout=3000, force=True)
                                            print(f"      ✅ Clicked '{button_text}' button: {field_title}")
                                        else:
                                            print(f"      ⚠️  '{button_text}' button not found in parent container")
                                    except:
                                        # Fallback: use global button index
                                        print(f"         ⚠️  Parent strategy failed, using global index")
                                        buttons = page.locator(f'button:has-text("{button_text}")').all()
                                        print(f"         📍 Found {len(buttons)} '{button_text}' buttons total")
                                        
                                        if button_text == "Yes":
                                            button_index = boolean_yes_clicks
                                            boolean_yes_clicks += 1
                                        else:
                                            button_index = boolean_no_clicks
                                            boolean_no_clicks += 1
                                            
                                        if button_index < len(buttons):
                                            btn = buttons[button_index]
                                            btn.scroll_into_view_if_needed()
                                            btn.click(timeout=3000, force=True)
                                            print(f"      ✅ Clicked '{button_text}' button #{button_index + 1}: {field_title}")
                                        
                                except Exception as e:
                                    print(f"      ⚠️  Boolean error: {str(e)[:80]}")
                            elif field_type == 'Location':
                                # Handle location autocomplete field
                                try:
                                    print(f"         🔍 Location field found, preparing to type...")
                                    element.scroll_into_view_if_needed()
                                    
                                    # Focus the field
                                    element.click()
                                    print(f"         ✓ Clicked location field")
                                    
                                    # Clear field first
                                    element.fill('')
                                    print(f"         ✓ Cleared field")
                                    
                                    # Type location character by character for reliability
                                    location_text = str(value)
                                    print(f"         📍 Typing location: '{location_text}'")
                                    
                                    # Try multiple typing methods
                                    try:
                                        # Method 1: press_sequentially (most reliable for autocomplete)
                                        element.press_sequentially(location_text, delay=random.randint(80, 120))
                                        
                                        # Verify typing worked
                                        filled = element.input_value()
                                        print(f"         ✓ After press_sequentially: '{filled}' ({len(filled) if filled else 0} chars)")
                                        
                                        if not filled or len(filled) < 3:
                                            # Method 2: Fallback to type()
                                            print(f"         ⚠️  press_sequentially didn't work, trying type()...")
                                            element.fill('')
                                            element.type(location_text, delay=100)
                                            filled = element.input_value()
                                            print(f"         ✓ After type(): '{filled}' ({len(filled) if filled else 0} chars)")
                                        
                                        if not filled or len(filled) < 3:
                                            # Method 3: Last resort - fill()
                                            print(f"         ⚠️  type() didn't work, trying fill()...")
                                            element.fill(location_text)
                                            filled = element.input_value()
                                            print(f"         ✓ After fill(): '{filled}' ({len(filled) if filled else 0} chars)")
                                    
                                    except Exception as type_error:
                                        print(f"         ⚠️  Typing error: {str(type_error)[:100]}")
                                        # Last resort
                                        element.fill(location_text)
                                    
                                    # Wait for autocomplete dropdown to appear
                                    page.wait_for_timeout(800)
                                    
                                    # Click the first autocomplete option if available
                                    try:
                                        autocomplete_option = page.locator('[role="option"]').first
                                        if autocomplete_option.count() > 0:
                                            autocomplete_option.click()
                                            print(f"      ✅ Selected from autocomplete: {field_title}")
                                        else:
                                            # Verify text was typed
                                            filled = element.input_value()
                                            if filled and len(filled) > 0:
                                                print(f"      ✅ Typed location: {field_title} ({len(filled)}/{len(location_text)} chars)")
                                            else:
                                                print(f"      ⚠️  Location field still empty after all attempts!")
                                    except Exception as e:
                                        filled = element.input_value()
                                        if filled:
                                            print(f"      ✅ Typed location: {field_title} ({len(filled)} chars)")
                                        else:
                                            print(f"      ⚠️  Location field empty!")
                                    
                                except Exception as e:
                                    print(f"      ⚠️  Location field error: {str(e)[:50]}")
                            elif field_type in ['ValueSelect', 'MultiValueSelect']:
                                # These might be rendered as dropdowns OR radio buttons OR checkboxes
                                # Check the actual HTML structure
                                try:
                                    print(f"         🔍 DEBUG: Processing {field_type} field: {field_title}")
                                    element.scroll_into_view_if_needed()
                                    
                                    # Get the tag name to understand what element we have
                                    tag_name = element.evaluate('el => el.tagName')
                                    elem_type = element.evaluate('el => el.type') if tag_name.lower() == 'input' else None
                                    print(f"         📄 Element: <{tag_name}> type={elem_type}")
                                    
                                    # Check if this element itself is a radio or checkbox
                                    if elem_type == 'radio':
                                        # This field IS a radio button (first one in group)
                                        print(f"         � This is a radio button input")
                                        
                                        # Find parent fieldset to get all options
                                        field_container = element.locator('xpath=ancestor::fieldset').first
                                        if field_container.count() == 0:
                                            field_container = element.locator('xpath=ancestor::div[contains(@class, "fieldEntry")]').first
                                        
                                        if field_container.count() > 0:
                                            # List all available options
                                            all_labels = field_container.locator('label').all()
                                            print(f"         � Available radio options ({len(all_labels)}):")
                                            for idx, lbl in enumerate(all_labels):
                                                lbl_text = lbl.inner_text()
                                                print(f"            [{idx}] '{lbl_text}'")
                                            
                                            # Find and click the matching label
                                            value_str = str(value)
                                            print(f"         🎯 Looking for: '{value_str}'")
                                            
                                            found = False
                                            for lbl in all_labels:
                                                lbl_text = lbl.inner_text().strip()
                                                # Case-insensitive comparison
                                                if lbl_text.lower() == value_str.lower():
                                                    print(f"         ✓ Found match: '{lbl_text}' (clicking...)")
                                                    lbl.click()
                                                    print(f"      ✅ Selected radio: {value_str}")
                                                    found = True
                                                    break
                                            
                                            if not found:
                                                print(f"      ⚠️  Could not find radio option: {value_str}")
                                                print(f"      🔍 Available were: {[l.inner_text().strip() for l in all_labels]}")
                                        else:
                                            print(f"      ⚠️  Could not find container for radio buttons")
                                    
                                    elif elem_type == 'checkbox':
                                        # This field IS a checkbox (first one in group)
                                        print(f"         ☑️  This is a checkbox input")
                                        
                                        # Find parent fieldset to get all checkboxes
                                        field_container = element.locator('xpath=ancestor::fieldset').first
                                        if field_container.count() == 0:
                                            field_container = element.locator('xpath=ancestor::div[contains(@class, "fieldEntry")]').first
                                        
                                        if field_container.count() > 0:
                                            # List all available checkboxes
                                            all_labels = field_container.locator('label').all()
                                            print(f"         📋 Available checkbox options ({len(all_labels)}):")
                                            for idx, lbl in enumerate(all_labels):
                                                lbl_text = lbl.inner_text()
                                                print(f"            [{idx}] '{lbl_text}'")
                                            
                                            if isinstance(value, list):
                                                values_to_select = value
                                            else:
                                                values_to_select = [value]
                                            
                                            print(f"         🎯 Need to select: {values_to_select}")
                                            
                                            # First, uncheck all checkboxes
                                            all_checkboxes = field_container.locator('input[type="checkbox"]').all()
                                            print(f"         🔄 Unchecking all {len(all_checkboxes)} checkboxes...")
                                            for idx, cb in enumerate(all_checkboxes):
                                                if cb.is_checked():
                                                    print(f"            [{idx}] Unchecking...")
                                                    cb.uncheck()
                                            
                                            # Now check the ones we want
                                            for v in values_to_select:
                                                print(f"         🔍 Looking for: '{v}'")
                                                found = False
                                                for lbl in all_labels:
                                                    lbl_text = lbl.inner_text().strip()
                                                    # Case-insensitive comparison
                                                    if lbl_text.lower() == v.lower():
                                                        print(f"         ✓ Found match: '{lbl_text}' (clicking...)")
                                                        lbl.click()
                                                        found = True
                                                        break
                                                if found:
                                                    print(f"         ✅ Checked: {v}")
                                                else:
                                                    print(f"         ⚠️  Not found: {v}")
                                                    print(f"         🔍 Available options were: {[l.inner_text().strip() for l in all_labels]}")
                                            
                                            print(f"      ✅ Selected checkboxes: {', '.join(values_to_select)}")
                                        else:
                                            print(f"      ⚠️  Could not find container for checkboxes")
                                    
                                    else:
                                        # Regular dropdown input
                                        print(f"         🔽 This is a dropdown/combobox")
                                        element.click()
                                        print(f"         ✓ Clicked dropdown")
                                        # Regular dropdown - original logic
                                        element.click()
                                        print(f"         🔍 Clicked dropdown for: {field_title}")
                                        
                                        if isinstance(value, list):
                                            # Multi-select dropdown
                                            for v in value:
                                                print(f"         📝 Typing to filter: {v}")
                                                element.type(str(v), delay=50)
                                                # Try to click the matching option
                                                try:
                                                    option = page.locator(f'[role="option"]:has-text("{v}")').first
                                                    if option.count() > 0:
                                                        option.click()
                                                        print(f"         ✓ Selected: {v}")
                                                except:
                                                    print(f"         ⚠️  Could not select: {v}")
                                        else:
                                            # Single select dropdown
                                            value_str = str(value)
                                            print(f"         📝 Typing to filter: {value_str}")
                                            element.type(value_str, delay=50)
                                            
                                            # Try clicking the matching option
                                            try:
                                                # Wait for options to appear
                                                page.wait_for_selector('[role="option"]', timeout=2000)
                                                
                                                # Try exact text match first
                                                option = page.locator(f'[role="option"]:has-text("{value_str}")').first
                                                if option.count() > 0:
                                                    option.click()
                                                    print(f"      ✅ Selected: {value_str}")
                                                else:
                                                    # Try pressing Enter to select first filtered option
                                                    element.press('Enter')
                                                    print(f"      ✅ Selected first option (pressed Enter)")
                                            except Exception as e:
                                                # Last resort: click first visible option
                                                try:
                                                    option = page.locator('[role="option"]').first
                                                    if option.count() > 0:
                                                        option.click()
                                                        print(f"      ✅ Selected first visible option")
                                                    else:
                                                        print(f"      ⚠️  No options found for: {value_str}")
                                                except:
                                                    print(f"      ⚠️  Could not select any option")
                                    
                                    time.sleep(random.uniform(0.3, 0.6))
                                except Exception as e:
                                    print(f"      ⚠️  Dropdown error: {str(e)[:80]}")
                            else:
                                # Handle text input with human-like typing
                                element.scroll_into_view_if_needed()
                                time.sleep(random.uniform(0.2, 0.5))
                                element.click()
                                time.sleep(random.uniform(0.3, 0.6))
                                
                                # Clear field first
                                element.fill('')
                                time.sleep(0.1)
                                
                                # Type character by character with delays for long text
                                text_to_type = str(value)
                                
                                # For longer text, use fill() then verify, for shorter use type()
                                if len(text_to_type) > 200:
                                    # For long text (like essays), use fill and verify
                                    element.fill(text_to_type)
                                    time.sleep(random.uniform(0.5, 1))
                                    
                                    # Verify it filled completely
                                    filled_value = element.input_value()
                                    if len(filled_value) < len(text_to_type) - 10:
                                        print(f"      ⚠️  Only filled {len(filled_value)}/{len(text_to_type)} chars, retrying...")
                                        element.fill('')
                                        time.sleep(0.2)
                                        # Try typing in chunks
                                        chunk_size = 100
                                        for i in range(0, len(text_to_type), chunk_size):
                                            chunk = text_to_type[i:i+chunk_size]
                                            element.type(chunk, delay=5)
                                            time.sleep(0.1)
                                else:
                                    # For shorter text, type with human-like delays
                                    element.type(text_to_type, delay=random.randint(20, 50))
                                
                                time.sleep(random.uniform(0.3, 0.7))
                                
                                # Final verification
                                filled_value = element.input_value()
                                expected_length = len(text_to_type)
                                actual_length = len(filled_value) if filled_value else 0
                                
                                if actual_length >= expected_length - 10:  # Allow small difference
                                    print(f"      ✅ Filled: {field_title} ({actual_length}/{expected_length} chars)")
                                else:
                                    print(f"      ⚠️  Incomplete fill: {field_title} ({actual_length}/{expected_length} chars)")
                                    print(f"         Trying one more time with press_sequentially...")
                                    element.click()
                                    element.press_sequentially(text_to_type, delay=10)
                                    time.sleep(0.5)
                            
                            # Human-like delay between fields
                            time.sleep(random.uniform(0.5, 1.5))
                        else:
                            print(f"      ⚠️  Could not find: {field_title}")
                    
                    except Exception as e:
                        print(f"      ⚠️  Error with {field_title}: {str(e)[:100]}")
                
                # Handle additional EEO fields that might not be in scraped data
                print(f"\n   🔍 Looking for additional EEO fields (Gender, Race, Veteran)...")
                
                # Save page HTML for debugging
                page_html = page.content()
                with open('form_html_debug.html', 'w', encoding='utf-8') as f:
                    f.write(page_html)
                print(f"   📄 Saved page HTML for debugging: form_html_debug.html")
                
                demographics = self.user_config.get('demographics', {})
                
                # Try to find and fill Gender identity radio buttons
                try:
                    gender_heading = page.locator('text=/Gender identity/i').first
                    if gender_heading.count() > 0:
                        print(f"   📋 Found Gender identity field")
                        gender_value = demographics.get('gender', 'Man')
                        
                        # Scroll to make sure it's visible
                        gender_heading.scroll_into_view_if_needed()
                        
                        # The heading label is the first child of fieldset, get the parent fieldset
                        gender_fieldset = gender_heading.locator('xpath=..').first
                        
                        # Find ALL labels within this fieldset (not just the heading)
                        all_labels = gender_fieldset.locator('label').all()
                        print(f"         🔍 Found {len(all_labels)} total labels in fieldset")
                        
                        # Filter out the heading label and show only option labels
                        option_labels = []
                        for label in all_labels:
                            label_text = label.inner_text().strip()
                            # Skip the "Gender identity" heading
                            if label_text and label_text.lower() != 'gender identity':
                                option_labels.append(label)
                                print(f"            - '{label_text}'")
                        
                        print(f"         🎯 Looking for: '{gender_value}'")
                        
                        # Find and click (case-insensitive)
                        matched = False
                        for label in option_labels:
                            label_text = label.inner_text().strip()
                            if label_text.lower() == gender_value.lower():
                                print(f"         ✓ Found match, clicking label...")
                                label.click()
                                print(f"      ✅ Selected Gender: {gender_value}")
                                matched = True
                                break
                        
                        if not matched:
                            print(f"      ⚠️  Could not find matching gender option: '{gender_value}'")
                            print(f"      Available options: {[l.inner_text().strip() for l in option_labels]}")
                except Exception as e:
                    print(f"      ⚠️  Gender field error: {str(e)[:80]}")
                
                # Try to find and fill Race identity dropdown
                try:
                    race_label = page.locator('text=/Race identity/i').first
                    if race_label.count() > 0:
                        print(f"   📋 Found Race identity field")
                        race_value = demographics.get('race', 'White')
                        # Find the dropdown input near the label
                        race_input = page.locator('input[placeholder="Start typing..." i]').last
                        if race_input.count() > 0:
                            race_input.scroll_into_view_if_needed()
                            race_input.click()
                            race_input.press_sequentially(race_value, delay=80)
                            # Click first autocomplete option
                            option = page.locator('[role="option"]').first
                            if option.count() > 0:
                                option.click()
                                print(f"      ✅ Selected Race: {race_value}")
                except Exception as e:
                    print(f"      ⚠️  Race field error: {str(e)[:50]}")
                
                # Try to find and fill Veteran status radio buttons
                try:
                    veteran_label = page.locator('text=/Veteran status/i').first
                    if veteran_label.count() > 0:
                        print(f"   📋 Found Veteran status field")
                        veteran_value = demographics.get('veteran_status', 'No')
                        # Map to radio button text
                        if veteran_value.lower() == 'no':
                            veteran_text = 'I am not a Veteran'
                        elif veteran_value.lower() == 'yes':
                            veteran_text = 'Yes, I am a Veteran'
                        else:
                            veteran_text = 'Prefer not to disclose'
                        
                        # Click the radio button
                        veteran_radio = page.locator(f'label:has-text("{veteran_text}")').first
                        if veteran_radio.count() > 0:
                            veteran_radio.click()
                            print(f"      ✅ Selected Veteran: {veteran_text}")
                except Exception as e:
                    print(f"      ⚠️  Veteran field error: {str(e)[:50]}")
                
                # Take screenshot BEFORE submitting to verify all fields
                print(f"\n   📸 Taking pre-submit screenshot to verify fields...")
                
                # Take screenshot immediately without scrolling delays
                pre_submit_screenshot = f"pre_submit_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                page.screenshot(path=pre_submit_screenshot, full_page=True)
                print(f"   ✅ Pre-submit screenshot: {pre_submit_screenshot}")
                
                # Check if auto-submit is enabled
                if not AUTO_SUBMIT:
                    print(f"\n" + "="*70)
                    print(f"   🛑 AUTO-SUBMIT DISABLED (Spam Detection Bypass)")
                    print(f"   " + "="*70)
                    print(f"   The form has been filled automatically.")
                    print(f"   Please review the form and click Submit MANUALLY.")
                    print(f"   ")
                    print(f"   Why? Ashby detects automated submissions. By having a human")
                    print(f"   click the final Submit button, it looks like a human applied.")
                    print(f"   " + "="*70)
                    print(f"   ⏸️  Review the form, then click Submit in the browser")
                    print(f"   ⏸️  Press ENTER after you've submitted to close browser...")
                    print(f"   " + "="*70 + "\n")
                    input()
                    print(f"   ✅ Closing browser...")
                    browser_context.close()
                    if profile_dir and profile_dir.startswith('/tmp'):
                        shutil.rmtree(profile_dir, ignore_errors=True)
                    return {
                        'success': True,
                        'status': 'manual_submit',
                        'message': 'Form filled, human submitted manually'
                    }
                
                # Submit the application (AUTO MODE)
                print(f"\n   ⏳ Looking for submit button...")
                
                # Find and click submit button
                submit_selectors = [
                    'button:has-text("Submit Application")',
                    'button:has-text("Submit")',
                    'button:has-text("Apply")',
                    'input[type="submit"]',
                    'button[type="submit"]'
                ]
                
                submitted = False
                for selector in submit_selectors:
                    try:
                        count = page.locator(selector).count()
                        print(f"      🔍 Checking selector '{selector}': found {count}")
                        if count > 0:
                            print(f"   🚀 Clicking submit button: {selector}")
                            button = page.locator(selector).first
                            
                            # Scroll to button and click
                            button.scroll_into_view_if_needed()
                            button.click()
                            
                            submitted = True
                            break
                    except Exception as e:
                        print(f"      ⚠️  Error with selector '{selector}': {str(e)[:50]}")
                        continue
                
                if submitted:
                    # Wait for submission to complete
                    print(f"   ⏳ Waiting for submission to process...")
                    
                    # Check for CAPTCHA
                    if page.locator('iframe[title*="recaptcha" i], iframe[src*="captcha" i]').count() > 0:
                        print(f"\n" + "="*70)
                        print(f"   🤖 CAPTCHA DETECTED!")
                        print(f"   " + "="*70)
                        print(f"   Please solve the CAPTCHA manually in the browser window.")
                        print(f"   Press ENTER after solving the CAPTCHA...")
                        print(f"   " + "="*70 + "\n")
                        input()
                    
                    # Check for success/confirmation
                    page.wait_for_load_state('networkidle', timeout=30000)
                    
                    # Take screenshot after submission
                    screenshot_path = f"application_screenshot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
                    page.screenshot(path=screenshot_path)
                    print(f"   📸 Screenshot saved: {screenshot_path}")
                    
                    # Check if we see confirmation or error
                    page_content = page.content().lower()
                    if 'spam' in page_content or 'flagged' in page_content:
                        print(f"\n   ⚠️  WARNING: Application may have been flagged as spam!")
                        status = 'flagged'
                        success = False
                        message = 'Application flagged as spam'
                    elif 'thank you' in page_content or 'submitted' in page_content or 'received' in page_content:
                        print(f"\n   ✅ Application submitted successfully!")
                        status = 'submitted'
                        success = True
                        message = 'Application submitted successfully'
                    else:
                        print(f"\n   ⚠️  Submission status unclear - check browser window")
                        status = 'uncertain'
                        success = True
                        message = 'Submitted but confirmation unclear'
                else:
                    print(f"\n   ❌ Could not find submit button")
                    status = 'failed'
                    success = False
                    message = 'Submit button not found'
                
                # Keep browser open for inspection
                print(f"\n   ⏸️  Browser kept open for inspection...")
                print(f"   ℹ️  Status: {status}")
                print(f"   ℹ️  Message: {message}")
                print(f"   ℹ️  Check the browser window to verify submission")
                print(f"   ℹ️  Press Ctrl+C when done\n")
                
                import signal
                def handler(signum, frame):
                    print(f"\n   🛑 Closing browser...")
                    browser_context.close()
                    # Cleanup temp profile directory
                    if profile_dir:
                        shutil.rmtree(profile_dir, ignore_errors=True)
                    exit(0)
                
                signal.signal(signal.SIGINT, handler)
                
                # Wait indefinitely
                while True:
                    time.sleep(1)
                
        except ImportError:
            print(f"   ❌ Playwright not installed")
            print(f"   Run: pip install playwright && playwright install chromium")
            return {
                'status': 'error',
                'success': False,
                'message': 'Playwright not installed'
            }
        except Exception as e:
            print(f"   ❌ Error: {e}")
            return {
                'status': 'error',
                'success': False,
                'message': str(e)
            }


def main():
    """Demo: Apply to Flagler Health job with auto-submission."""
    
    # Initialize AI Job Bidder
    bidder = AIJobBidder('config.json')
    
    # Job URL
    job_url = "https://jobs.ashbyhq.com/rula/1850600a-6c12-413a-b008-ee442f01a592"
    
    print(f"\n{'='*70}")
    print("🤖 AI JOB BIDDER - AUTO APPLICATION SYSTEM")
    print(f"{'='*70}\n")
    print(f"Target: {job_url}\n")
    
    # Step 1: Extract application form
    print("Step 1/4: Extracting application form...")
    form_data = bidder.extract_application_form(job_url)
    
    if not form_data:
        print("❌ Failed to extract application form")
        return
    
    # Step 2: Fill application with AI assistance
    print("\nStep 2/4: Filling application with AI...")
    filled_application = bidder.fill_application(form_data)
    
    # Step 3: Submit application with browser automation
    print("\nStep 3/4: Submitting application with browser...")
    submission_result = bidder.submit_application_with_browser(
        filled_application, 
        form_data,
        headless=False  # Set to True to hide browser window
    )
    
    # Step 4: Save application data
    print("Step 4/4: Saving application data...")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f'filled_application_{timestamp}.json'
    
    # Add submission result to data
    filled_application['submission_result'] = submission_result
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(filled_application, f, indent=2, ensure_ascii=False)
    
    print(f"   💾 Application saved to: {output_file}")
    
    # Generate preview
    preview = bidder.generate_application_preview(filled_application)
    print(f"\n{preview}")
    
    # Final summary
    print(f"\n{'='*70}")
    print("📊 SUBMISSION SUMMARY")
    print(f"{'='*70}")
    print(f"Status: {submission_result['status'].upper()}")
    print(f"Success: {'✅ Yes' if submission_result['success'] else '❌ No'}")
    print(f"Message: {submission_result['message']}")
    
    if submission_result['status'] == 'dry_run':
        print(f"\n💡 To actually submit, change dry_run=False in main()")
    
    print(f"\n✅ AI-powered application complete!")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
