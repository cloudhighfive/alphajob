"""
Simple and reliable DOCX resume tailoring that preserves formatting.
Directly updates paragraph text content while keeping all formatting intact.
"""

from docx import Document
from typing import Dict, List
import re
import sys
import os

# Add parent directory to path to import from src
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from src.utils.docx_formatter import ResumeFormattingPreserver, DocxFormatter
from src.utils.logger import get_logger

logger = get_logger(__name__)


def extract_resume_content(doc: Document) -> Dict:
    """
    Extract resume content by sections.
    Returns dict with section start/end indices and original text.
    """
    content = {
        'summary': {'start': -1, 'end': -1, 'paras': []},
        'skills': {'start': -1, 'end': -1, 'paras': []},
        'experience_jobs': []  # List of {start, end, company, paras}
    }
    
    current_section = None
    current_job = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_lower = text.lower()
        
        if not text:
            continue
        
        # Detect section headings
        if 'summary of qualifications' in text_lower or text_lower == 'summary':
            current_section = 'summary'
            content['summary']['start'] = i + 1  # Content starts after heading
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'technical skills' in text_lower:
            current_section = 'skills'
            content['skills']['start'] = i + 1
            if content['summary']['end'] == -1 and content['summary']['start'] > 0:
                content['summary']['end'] = i - 1
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'relevant work experience' in text_lower or text_lower == 'experience':
            current_section = 'experience'
            if content['skills']['end'] == -1 and content['skills']['start'] > 0:
                content['skills']['end'] = i - 1
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'education' in text_lower or 'certifications' in text_lower:
            current_section = 'other'
            if content['skills']['end'] == -1 and content['skills']['start'] > 0:
                content['skills']['end'] = i - 1
            if current_job:
                current_job['end'] = i - 1
                content['experience_jobs'].append(current_job)
                current_job = None
        
        # Handle content based on current section
        elif current_section == 'summary':
            content['summary']['paras'].append({'index': i, 'text': text})
        
        elif current_section == 'skills':
            content['skills']['paras'].append({'index': i, 'text': text})
        
        elif current_section == 'experience':
            # Check if new job entry (has tab for company/location or date range)
            if '\t' in text or (any(month in text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']) and '–' in text):
                if current_job and 'bullets_start' in current_job:
                    # Close previous job
                    current_job['end'] = i - 1
                    content['experience_jobs'].append(current_job)
                
                # Start new job or continue job header
                if current_job is None or 'dates' in current_job:
                    current_job = {
                        'start': i,
                        'company': '',
                        'location': '',
                        'title': '',
                        'dates': '',
                        'paras': []
                    }
                
                # Parse company/location or title/dates
                if '–' in text and any(m in text for m in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                    parts = text.split('\t')
                    current_job['title'] = parts[0].strip()
                    current_job['dates'] = parts[-1].strip() if len(parts) > 1 else text
                    current_job['bullets_start'] = i + 1
                else:
                    parts = text.split('\t')
                    current_job['company'] = parts[0].strip()
                    current_job['location'] = parts[-1].strip() if len(parts) > 1 else ''
            
            elif current_job and 'bullets_start' in current_job:
                # This is a bullet point
                current_job['paras'].append({'index': i, 'text': text})
    
    # Close last job
    if current_job:
        current_job['end'] = len(doc.paragraphs) - 1
        content['experience_jobs'].append(current_job)
    
    # Set final ends
    if content['summary']['end'] == -1 and content['summary']['start'] > 0:
        content['summary']['end'] = content['summary']['paras'][-1]['index'] if content['summary']['paras'] else content['summary']['start']
    
    if content['skills']['end'] == -1 and content['skills']['start'] > 0:
        content['skills']['end'] = content['skills']['paras'][-1]['index'] if content['skills']['paras'] else content['skills']['start']
    
    return content


def clean_ai_response(text: str) -> str:
    """
    Remove AI commentary and unwanted formatting from responses.
    """
    # Common AI prefixes/patterns to remove (case-insensitive)
    patterns_to_remove = [
        r'^Here is the rewritten.*?:\s*',
        r'^Here are the rewritten.*?:\s*',
        r'^Here\'s the rewritten.*?:\s*',
        r'^Here are the.*?:\s*',
        r'^Here is the.*?:\s*',
        r'^Here\'s the.*?:\s*',
        r'^The rewritten.*?:\s*',
        r'^Rewritten.*?:\s*',
        r'^Tailored.*?:\s*',
        r'^Here is a tailored.*?:\s*',
        r'^Here are tailored.*?:\s*',
        r'^I\'ve tailored.*?:\s*',
        r'^I have tailored.*?:\s*',
    ]
    
    # Remove AI commentary patterns
    import re
    cleaned = text
    for pattern in patterns_to_remove:
        cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE | re.MULTILINE)
    
    # Split into lines and process
    lines = cleaned.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip empty lines
        if not line_stripped:
            continue
        
        # Skip lines that are pure commentary (common patterns)
        skip_patterns = [
            r'^(Here|I\'ve|I have|The following|These are|This is)\s',
            r'^(Note|Important|Remember):',
        ]
        
        skip = False
        for pattern in skip_patterns:
            if re.match(pattern, line_stripped, re.IGNORECASE):
                skip = True
                break
        
        if skip:
            continue
        
        # Remove bullet point symbols at start (we'll add them back in resume)
        cleaned = line_stripped
        
        # Remove bullet point symbols
        if cleaned.startswith('• '):
            cleaned = cleaned[2:]
        elif cleaned.startswith('- '):
            cleaned = cleaned[2:]
        elif cleaned.startswith('* '):
            cleaned = cleaned[2:]
        
        # Remove numbered list markers (1. 2. 3. etc.)
        cleaned = re.sub(r'^\d+\.\s+', '', cleaned)
        
        # Remove markdown bold markers
        cleaned = cleaned.replace('**', '')
        
        if cleaned:
            cleaned_lines.append(cleaned)
    
    return '\n'.join(cleaned_lines)


def update_resume_sections(doc: Document, content: Dict, tailored_summary: str, 
                          tailored_skills: str, tailored_jobs: List[Dict]) -> Document:
    """
    Update resume sections with tailored content.
    Modifies doc in place and returns it.
    Preserves original formatting including font, size, style, alignment, and spacing.
    """
    preserver = ResumeFormattingPreserver()
    
    # Clean AI responses
    tailored_summary = clean_ai_response(tailored_summary)
    tailored_skills = clean_ai_response(tailored_skills)
    
    logger.info(f"\n{'='*70}")
    logger.info(f"📝 APPLYING TAILORED CONTENT TO DOCUMENT")
    logger.info(f"{'='*70}\n")
    
    # Update summary paragraphs - preserve formatting
    if content['summary']['paras'] and tailored_summary:
        logger.info("🔄 Updating SUMMARY section...")
        summary_lines = [l.strip() for l in tailored_summary.split('\n') if l.strip()]
        
        for i, para_info in enumerate(content['summary']['paras']):
            if i < len(summary_lines):
                para = doc.paragraphs[para_info['index']]
                # Use advanced formatter to preserve all formatting
                formatter = DocxFormatter()
                formatter.update_paragraph_text_preserve_format(para, summary_lines[i])
                logger.debug(f"   ✅ Updated summary para {i}")
        
        logger.info(f"   ✅ Summary section complete ({len(summary_lines)} lines)\n")
    
    # Update skills paragraphs - preserve formatting
    if content['skills']['paras'] and tailored_skills:
        logger.info("🔄 Updating TECHNICAL SKILLS section...")
        skills_lines = [l.strip() for l in tailored_skills.split('\n') if l.strip()]
        
        logger.info(f"   📊 Original paragraphs: {len(content['skills']['paras'])}")
        logger.info(f"   📊 Tailored lines: {len(skills_lines)}")
        
        # Update each paragraph with corresponding line
        for i, para_info in enumerate(content['skills']['paras']):
            para = doc.paragraphs[para_info['index']]
            
            if i < len(skills_lines):
                # Use advanced formatter to preserve all formatting
                formatter = DocxFormatter()
                formatter.update_paragraph_text_preserve_format(para, skills_lines[i])
                logger.debug(f"   ✅ Updated skills para {i}: {skills_lines[i][:60]}...")
            else:
                logger.warning(f"   ⚠️  No tailored content for paragraph {i}, keeping original")
        
        logger.info(f"   ✅ Skills section complete\n")
    
    # Update experience bullets for each job - preserve formatting
    if tailored_jobs and content['experience_jobs']:
        logger.info(f"🔄 Updating WORK EXPERIENCE ({len(content['experience_jobs'])} jobs)...")
        
        for job_idx, job in enumerate(content['experience_jobs']):
            if job_idx >= len(tailored_jobs):
                logger.warning(f"   ⚠️  No tailored content for job {job_idx+1}")
                continue
            
            if not job['paras']:
                logger.warning(f"   ⚠️  Job {job_idx+1} has no original paragraphs")
                continue
            
            tailored_bullets = tailored_jobs[job_idx].get('bullets', [])
            # Clean each bullet
            tailored_bullets = [clean_ai_response(b).strip() for b in tailored_bullets if clean_ai_response(b).strip()]
            
            logger.info(f"\n   📝 Job {job_idx+1}/{len(content['experience_jobs'])}: {job.get('company', 'Unknown')}")
            logger.info(f"      Original bullets: {len(job['paras'])}")
            logger.info(f"      Tailored bullets: {len(tailored_bullets)}")
            
            for i, para_info in enumerate(job['paras']):
                if i >= len(tailored_bullets):
                    logger.warning(f"      ⚠️  No tailored bullet for position {i}, keeping original")
                    continue
                
                para = doc.paragraphs[para_info['index']]
                # Use advanced formatter to preserve all formatting
                formatter = DocxFormatter()
                formatter.update_paragraph_text_preserve_format(para, tailored_bullets[i])
                logger.debug(f"      ✅ Updated bullet {i+1}")
            
            logger.info(f"      ✅ Job {job_idx+1} complete")
        
        logger.info(f"\n   ✅ All experience sections updated\n")
    
    logger.info(f"{'='*70}")
    logger.info(f"✅ DOCUMENT UPDATE COMPLETE")
    logger.info(f"{'='*70}\n")
    
    return doc
