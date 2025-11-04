"""
Simple and reliable DOCX resume tailoring that preserves formatting.
Directly updates paragraph text content while keeping all formatting intact.
"""

from docx import Document
from typing import Dict, List
import re


def extract_resume_content(doc: Document) -> Dict:
    """
    Extract resume content by sections.
    Returns dict with section start/end indices and original text.
    """
    content = {
        'summary': {'start': -1, 'end': -1, 'paras': []},
        'skills': {'start': -1, 'end': -1, 'paras': []},
        'experience_jobs': []  # List of {start, end, company, paras}
    }
    
    current_section = None
    current_job = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        text_lower = text.lower()
        
        if not text:
            continue
        
        # Detect section headings
        if 'summary of qualifications' in text_lower or text_lower == 'summary':
            current_section = 'summary'
            content['summary']['start'] = i + 1  # Content starts after heading
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'technical skills' in text_lower:
            current_section = 'skills'
            content['skills']['start'] = i + 1
            if content['summary']['end'] == -1 and content['summary']['start'] > 0:
                content['summary']['end'] = i - 1
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'relevant work experience' in text_lower or text_lower == 'experience':
            current_section = 'experience'
            if content['skills']['end'] == -1 and content['skills']['start'] > 0:
                content['skills']['end'] = i - 1
            if current_job:
                content['experience_jobs'].append(current_job)
                current_job = None
        
        elif 'education' in text_lower or 'certifications' in text_lower:
            current_section = 'other'
            if content['skills']['end'] == -1 and content['skills']['start'] > 0:
                content['skills']['end'] = i - 1
            if current_job:
                current_job['end'] = i - 1
                content['experience_jobs'].append(current_job)
                current_job = None
        
        # Handle content based on current section
        elif current_section == 'summary':
            content['summary']['paras'].append({'index': i, 'text': text})
        
        elif current_section == 'skills':
            content['skills']['paras'].append({'index': i, 'text': text})
        
        elif current_section == 'experience':
            # Check if new job entry (has tab for company/location or date range)
            if '\t' in text or (any(month in text for month in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']) and '–' in text):
                if current_job and 'bullets_start' in current_job:
                    # Close previous job
                    current_job['end'] = i - 1
                    content['experience_jobs'].append(current_job)
                
                # Start new job or continue job header
                if current_job is None or 'dates' in current_job:
                    current_job = {
                        'start': i,
                        'company': '',
                        'location': '',
                        'title': '',
                        'dates': '',
                        'paras': []
                    }
                
                # Parse company/location or title/dates
                if '–' in text and any(m in text for m in ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']):
                    parts = text.split('\t')
                    current_job['title'] = parts[0].strip()
                    current_job['dates'] = parts[-1].strip() if len(parts) > 1 else text
                    current_job['bullets_start'] = i + 1
                else:
                    parts = text.split('\t')
                    current_job['company'] = parts[0].strip()
                    current_job['location'] = parts[-1].strip() if len(parts) > 1 else ''
            
            elif current_job and 'bullets_start' in current_job:
                # This is a bullet point
                current_job['paras'].append({'index': i, 'text': text})
    
    # Close last job
    if current_job:
        current_job['end'] = len(doc.paragraphs) - 1
        content['experience_jobs'].append(current_job)
    
    # Set final ends
    if content['summary']['end'] == -1 and content['summary']['start'] > 0:
        content['summary']['end'] = content['summary']['paras'][-1]['index'] if content['summary']['paras'] else content['summary']['start']
    
    if content['skills']['end'] == -1 and content['skills']['start'] > 0:
        content['skills']['end'] = content['skills']['paras'][-1]['index'] if content['skills']['paras'] else content['skills']['start']
    
    return content


def clean_ai_response(text: str) -> str:
    """
    Remove AI commentary and unwanted formatting from responses.
    """
    # Common AI prefixes to remove
    prefixes = [
        "Here is the rewritten",
        "Here are the rewritten",
        "Here's the rewritten",
        "Here are the",
        "Here is the",
        "The rewritten",
        "Rewritten"
    ]
    
    lines = text.split('\n')
    cleaned_lines = []
    
    for line in lines:
        line_stripped = line.strip()
        
        # Skip lines that are just AI commentary
        skip = False
        for prefix in prefixes:
            if line_stripped.lower().startswith(prefix.lower()):
                skip = True
                break
        
        if skip or not line_stripped:
            continue
        
        # Remove bullet points and numbered list markers
        # Remove: •, -, *, 1., 2., etc.
        cleaned = line_stripped
        
        # Remove bullet point symbols at start
        if cleaned.startswith('• '):
            cleaned = cleaned[2:]
        elif cleaned.startswith('- '):
            cleaned = cleaned[2:]
        elif cleaned.startswith('* '):
            cleaned = cleaned[2:]
        
        # Remove numbered list markers (1. 2. 3. etc.)
        import re
        cleaned = re.sub(r'^\d+\.\s+', '', cleaned)
        
        # Remove markdown bold markers
        cleaned = cleaned.replace('**', '')
        
        if cleaned:
            cleaned_lines.append(cleaned)
    
    return '\n'.join(cleaned_lines)


def update_resume_sections(doc: Document, content: Dict, tailored_summary: str, 
                          tailored_skills: str, tailored_jobs: List[Dict]) -> Document:
    """
    Update resume sections with tailored content.
    Modifies doc in place and returns it.
    Preserves original formatting including font, size, and style.
    """
    # Clean AI responses
    tailored_summary = clean_ai_response(tailored_summary)
    tailored_skills = clean_ai_response(tailored_skills)
    
    # Update summary paragraphs - preserve formatting
    if content['summary']['paras'] and tailored_summary:
        summary_lines = [l.strip() for l in tailored_summary.split('\n') if l.strip()]
        for i, para_info in enumerate(content['summary']['paras']):
            if i < len(summary_lines):
                para = doc.paragraphs[para_info['index']]
                # Store original formatting
                original_runs = list(para.runs)
                
                # Clear existing text
                for run in original_runs:
                    run.text = ''
                
                # Add new text to first run (preserves font formatting)
                if para.runs:
                    para.runs[0].text = summary_lines[i]
                else:
                    # No runs exist, create one with new text
                    para.add_run(summary_lines[i])
    
    # Update skills paragraphs - preserve formatting
    if content['skills']['paras'] and tailored_skills:
        skills_lines = [l.strip() for l in tailored_skills.split('\n') if l.strip()]
        for i, para_info in enumerate(content['skills']['paras']):
            if i < len(skills_lines):
                para = doc.paragraphs[para_info['index']]
                # Store original formatting
                original_runs = list(para.runs)
                
                # Clear existing text
                for run in original_runs:
                    run.text = ''
                
                # Add new text to first run
                if para.runs:
                    para.runs[0].text = skills_lines[i]
                else:
                    para.add_run(skills_lines[i])
    
    # Update experience bullets for each job - preserve formatting
    for job_idx, job in enumerate(content['experience_jobs']):
        if job_idx < len(tailored_jobs) and job['paras']:
            tailored_bullets = tailored_jobs[job_idx].get('bullets', [])
            # Clean each bullet
            tailored_bullets = [clean_ai_response(b).strip() for b in tailored_bullets if clean_ai_response(b).strip()]
            
            for i, para_info in enumerate(job['paras']):
                if i < len(tailored_bullets):
                    para = doc.paragraphs[para_info['index']]
                    # Store original formatting
                    original_runs = list(para.runs)
                    
                    # Clear existing text
                    for run in original_runs:
                        run.text = ''
                    
                    # Add new text to first run
                    if para.runs:
                        para.runs[0].text = tailored_bullets[i]
                    else:
                        para.add_run(tailored_bullets[i])
    
    return doc
