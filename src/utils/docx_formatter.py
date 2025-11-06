"""
Advanced DOCX formatting preservation utilities.
Handles font families, sizes, colors, bold, italic, alignment, spacing, etc.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
import copy

from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocxFormatter:
    """Handle advanced DOCX formatting preservation."""
    
    @staticmethod
    def capture_paragraph_format(paragraph) -> Dict[str, Any]:
        """
        Capture all formatting details from a paragraph.
        
        Returns dict with:
        - alignment
        - spacing (before, after, line_spacing)
        - indentation (left, right, first_line)
        - runs: list of run formats (font, size, bold, italic, color, etc.)
        """
        format_info = {
            'alignment': paragraph.alignment,
            'spacing': {
                'before': paragraph.paragraph_format.space_before,
                'after': paragraph.paragraph_format.space_after,
                'line_spacing': paragraph.paragraph_format.line_spacing,
                'line_spacing_rule': paragraph.paragraph_format.line_spacing_rule
            },
            'indentation': {
                'left': paragraph.paragraph_format.left_indent,
                'right': paragraph.paragraph_format.right_indent,
                'first_line': paragraph.paragraph_format.first_line_indent
            },
            'runs': []
        }
        
        # Capture run formatting (fonts, sizes, styles)
        for run in paragraph.runs:
            run_format = {
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color.rgb if run.font.color.rgb else None,
                'highlight_color': run.font.highlight_color
            }
            format_info['runs'].append(run_format)
        
        return format_info
    
    @staticmethod
    def apply_paragraph_format(paragraph, format_info: Dict[str, Any]):
        """
        Apply captured formatting to a paragraph.
        
        Args:
            paragraph: Target paragraph
            format_info: Formatting information captured by capture_paragraph_format
        """
        # Apply paragraph-level formatting
        if format_info.get('alignment'):
            paragraph.alignment = format_info['alignment']
        
        # Apply spacing
        spacing = format_info.get('spacing', {})
        if spacing.get('before'):
            paragraph.paragraph_format.space_before = spacing['before']
        if spacing.get('after'):
            paragraph.paragraph_format.space_after = spacing['after']
        if spacing.get('line_spacing'):
            paragraph.paragraph_format.line_spacing = spacing['line_spacing']
        if spacing.get('line_spacing_rule'):
            paragraph.paragraph_format.line_spacing_rule = spacing['line_spacing_rule']
        
        # Apply indentation
        indentation = format_info.get('indentation', {})
        if indentation.get('left'):
            paragraph.paragraph_format.left_indent = indentation['left']
        if indentation.get('right'):
            paragraph.paragraph_format.right_indent = indentation['right']
        if indentation.get('first_line'):
            paragraph.paragraph_format.first_line_indent = indentation['first_line']
    
    @staticmethod
    def apply_run_format(run, run_format: Dict[str, Any]):
        """
        Apply formatting to a run.
        
        Args:
            run: Target run
            run_format: Run formatting info
        """
        if run_format.get('font_name'):
            run.font.name = run_format['font_name']
        if run_format.get('font_size'):
            run.font.size = run_format['font_size']
        if run_format.get('bold') is not None:
            run.font.bold = run_format['bold']
        if run_format.get('italic') is not None:
            run.font.italic = run_format['italic']
        if run_format.get('underline') is not None:
            run.font.underline = run_format['underline']
        if run_format.get('color'):
            run.font.color.rgb = run_format['color']
        if run_format.get('highlight_color'):
            run.font.highlight_color = run_format['highlight_color']
    
    @staticmethod
    def update_paragraph_text_preserve_format(paragraph, new_text: str):
        """
        Update paragraph text while preserving all formatting.
        
        Strategy:
        1. Capture original formatting from first run
        2. Clear all runs
        3. Create new run with new text
        4. Apply captured formatting
        
        Args:
            paragraph: Paragraph to update
            new_text: New text content
        """
        # Capture formatting from first run (if exists)
        original_format = None
        if paragraph.runs:
            original_format = {
                'font_name': paragraph.runs[0].font.name,
                'font_size': paragraph.runs[0].font.size,
                'bold': paragraph.runs[0].font.bold,
                'italic': paragraph.runs[0].font.italic,
                'underline': paragraph.runs[0].font.underline,
                'color': paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color.rgb else None
            }
        
        # Clear all runs
        for run in list(paragraph.runs):
            run.text = ''
        
        # Add new text in first run or create new run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            target_run = paragraph.runs[0]
        else:
            target_run = paragraph.add_run(new_text)
        
        # Apply original formatting
        if original_format:
            DocxFormatter.apply_run_format(target_run, original_format)
    
    @staticmethod
    def clone_paragraph_format(source_para, target_para):
        """
        Clone all formatting from source paragraph to target paragraph.
        Does not change text content.
        
        Args:
            source_para: Source paragraph to copy formatting from
            target_para: Target paragraph to apply formatting to
        """
        # Clone paragraph format
        target_para.paragraph_format.alignment = source_para.paragraph_format.alignment
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
        target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
        
        # Clone run formats (apply source first run format to target first run)
        if source_para.runs and target_para.runs:
            source_run = source_para.runs[0]
            target_run = target_para.runs[0]
            
            target_run.font.name = source_run.font.name
            target_run.font.size = source_run.font.size
            target_run.font.bold = source_run.font.bold
            target_run.font.italic = source_run.font.italic
            target_run.font.underline = source_run.font.underline
            if source_run.font.color.rgb:
                target_run.font.color.rgb = source_run.font.color.rgb


class ResumeFormattingPreserver:
    """Specialized formatter for resume tailoring with format preservation."""
    
    def __init__(self):
        self.formatter = DocxFormatter()
    
    def update_section_preserve_format(
        self,
        doc: Document,
        section_paragraphs: List[Dict],
        new_content_lines: List[str]
    ):
        """
        Update a section's paragraphs with new content while preserving formatting.
        
        Args:
            doc: Document object
            section_paragraphs: List of dicts with 'index' and 'text'
            new_content_lines: List of new text lines
        """
        logger.info(f"   📝 Updating {len(section_paragraphs)} paragraphs with {len(new_content_lines)} new lines")
        
        for i, para_info in enumerate(section_paragraphs):
            if i >= len(new_content_lines):
                logger.warning(f"   ⚠️  No content for paragraph {i}, keeping original")
                continue
            
            para_idx = para_info['index']
            new_text = new_content_lines[i]
            
            if para_idx >= len(doc.paragraphs):
                logger.error(f"   ❌ Paragraph index {para_idx} out of range")
                continue
            
            para = doc.paragraphs[para_idx]
            
            # Capture original formatting
            original_format = self.formatter.capture_paragraph_format(para)
            
            # Update text using our formatter
            self.formatter.update_paragraph_text_preserve_format(para, new_text)
            
            # Re-apply paragraph-level formatting (alignment, spacing)
            self.formatter.apply_paragraph_format(para, original_format)
            
            logger.debug(f"   ✅ Updated para {i}: {new_text[:60]}...")
    
    def log_formatting_details(self, doc: Document, para_indices: List[int]):
        """
        Log detailed formatting information for debugging.
        
        Args:
            doc: Document object
            para_indices: List of paragraph indices to inspect
        """
        logger.info(f"\n{'='*70}")
        logger.info(f"🔍 FORMATTING DETAILS")
        logger.info(f"{'='*70}")
        
        for idx in para_indices[:5]:  # Limit to first 5 for brevity
            if idx >= len(doc.paragraphs):
                continue
            
            para = doc.paragraphs[idx]
            format_info = self.formatter.capture_paragraph_format(para)
            
            logger.info(f"\nParagraph {idx}: {para.text[:50]}...")
            logger.info(f"  Alignment: {format_info['alignment']}")
            logger.info(f"  Spacing: before={format_info['spacing']['before']}, after={format_info['spacing']['after']}")
            
            if format_info['runs']:
                run_fmt = format_info['runs'][0]
                logger.info(f"  Font: {run_fmt['font_name']}, Size: {run_fmt['font_size']}")
                logger.info(f"  Bold: {run_fmt['bold']}, Italic: {run_fmt['italic']}")
        
        logger.info(f"{'='*70}\n")
