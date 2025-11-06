"""
Resume quality validation and ATS optimization scoring.
"""

from docx import Document
from typing import Dict, List, Tuple, Set
import re
from collections import Counter

from src.utils.logger import get_logger

logger = get_logger(__name__)


class ResumeValidator:
    """Validate resume quality and ATS compatibility."""
    
    def __init__(self):
        self.min_keyword_density = 0.05  # 5%
        self.target_keyword_density = 0.10  # 10%
        self.min_bullet_length = 30  # characters
        self.max_bullet_length = 200  # characters
    
    def extract_keywords_from_job_description(self, job_description: str) -> Dict[str, int]:
        """
        Extract important keywords from job description.
        
        Returns:
            Dict mapping keyword to frequency in JD
        """
        # Common tech keywords and patterns
        tech_patterns = [
            # Languages
            r'\b(python|java|javascript|typescript|kotlin|go|rust|c\+\+|c#|ruby|php|swift)\b',
            # Frameworks
            r'\b(react|vue|angular|node\.?js|express|django|flask|spring|fastapi|next\.?js)\b',
            # Cloud/DevOps
            r'\b(aws|azure|gcp|docker|kubernetes|k8s|jenkins|terraform|ansible|ci/cd)\b',
            # Databases
            r'\b(postgresql|mysql|mongodb|redis|elasticsearch|dynamodb|sql)\b',
            # Tools
            r'\b(git|github|gitlab|jira|confluence|slack)\b',
            # AI/ML
            r'\b(tensorflow|pytorch|scikit-learn|pandas|numpy|llm|gpt|transformer|hugging\s*face)\b',
            # Architecture
            r'\b(microservices|rest|restful|api|graphql|grpc|serverless)\b',
            # Concepts
            r'\b(agile|scrum|tdd|ci/cd|devops|mlops|data\s*pipeline)\b'
        ]
        
        jd_lower = job_description.lower()
        keyword_counts = Counter()
        
        for pattern in tech_patterns:
            matches = re.findall(pattern, jd_lower, re.IGNORECASE)
            for match in matches:
                keyword_counts[match.lower()] += 1
        
        # Also extract repeated noun phrases (2-3 words)
        words = re.findall(r'\b[a-z]{3,}\b', jd_lower)
        word_counts = Counter(words)
        
        # Add high-frequency words (appearing 3+ times)
        for word, count in word_counts.items():
            if count >= 3 and len(word) > 3:
                if word not in keyword_counts:
                    keyword_counts[word] = count
        
        return dict(keyword_counts.most_common(50))
    
    def calculate_keyword_density(
        self,
        resume_text: str,
        keywords: Dict[str, int]
    ) -> Dict[str, Dict]:
        """
        Calculate keyword density in resume.
        
        Args:
            resume_text: Full resume text
            keywords: Dict of keywords from job description
            
        Returns:
            Dict with keyword stats
        """
        resume_lower = resume_text.lower()
        total_words = len(resume_text.split())
        
        keyword_stats = {}
        total_keyword_matches = 0
        
        for keyword, jd_count in keywords.items():
            resume_count = resume_lower.count(keyword.lower())
            if resume_count > 0:
                coverage = (resume_count / jd_count * 100) if jd_count > 0 else 0
                keyword_stats[keyword] = {
                    'jd_count': jd_count,
                    'resume_count': resume_count,
                    'coverage': coverage,
                    'status': 'excellent' if coverage >= 50 else 'good' if coverage >= 30 else 'low'
                }
                total_keyword_matches += resume_count
        
        overall_density = (total_keyword_matches / total_words) if total_words > 0 else 0
        
        return {
            'keywords': keyword_stats,
            'overall_density': overall_density,
            'total_matches': total_keyword_matches,
            'total_words': total_words
        }
    
    def validate_resume_structure(self, doc: Document) -> Dict[str, any]:
        """
        Validate resume structure and sections.
        
        Returns:
            Dict with validation results
        """
        issues = []
        sections_found = {
            'summary': False,
            'skills': False,
            'experience': False,
            'education': False
        }
        
        bullet_count = 0
        short_bullets = []
        long_bullets = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            
            # Check for required sections
            if 'summary' in text or 'qualifications' in text:
                sections_found['summary'] = True
            elif 'skills' in text or 'technical' in text:
                sections_found['skills'] = True
            elif 'experience' in text or 'employment' in text:
                sections_found['experience'] = True
            elif 'education' in text:
                sections_found['education'] = True
            
            # Check bullet point length
            if len(text) > 20 and not any(keyword in text for keyword in ['summary', 'skills', 'experience', 'education']):
                bullet_count += 1
                if len(text) < self.min_bullet_length:
                    short_bullets.append((i, text[:50]))
                elif len(text) > self.max_bullet_length:
                    long_bullets.append((i, text[:50]))
        
        # Check for missing sections
        for section, found in sections_found.items():
            if not found:
                issues.append(f"Missing {section} section")
        
        if short_bullets:
            issues.append(f"{len(short_bullets)} bullets are too short (< {self.min_bullet_length} chars)")
        
        if long_bullets:
            issues.append(f"{len(long_bullets)} bullets are too long (> {self.max_bullet_length} chars)")
        
        return {
            'sections_found': sections_found,
            'bullet_count': bullet_count,
            'short_bullets': len(short_bullets),
            'long_bullets': len(long_bullets),
            'issues': issues,
            'valid': len(issues) == 0
        }
    
    def calculate_ats_score(
        self,
        doc: Document,
        job_description: str
    ) -> Dict[str, any]:
        """
        Calculate comprehensive ATS compatibility score.
        
        Returns:
            Dict with overall score and breakdown
        """
        logger.info(f"\n{'='*70}")
        logger.info(f"🎯 CALCULATING ATS SCORE")
        logger.info(f"{'='*70}\n")
        
        # Extract resume text
        resume_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        
        # 1. Structure validation (20 points)
        structure_result = self.validate_resume_structure(doc)
        structure_score = 20 if structure_result['valid'] else max(0, 20 - len(structure_result['issues']) * 4)
        
        logger.info(f"📋 Structure Score: {structure_score}/20")
        if structure_result['issues']:
            for issue in structure_result['issues']:
                logger.info(f"   ⚠️  {issue}")
        else:
            logger.info(f"   ✅ All required sections present")
        
        # 2. Keyword matching (50 points)
        keywords = self.extract_keywords_from_job_description(job_description)
        keyword_stats = self.calculate_keyword_density(resume_text, keywords)
        
        # Calculate keyword score based on coverage
        excellent_count = sum(1 for k, v in keyword_stats['keywords'].items() if v['status'] == 'excellent')
        good_count = sum(1 for k, v in keyword_stats['keywords'].items() if v['status'] == 'good')
        total_keywords = len(keywords)
        
        coverage_rate = ((excellent_count * 1.0 + good_count * 0.6) / total_keywords) if total_keywords > 0 else 0
        keyword_score = min(50, int(coverage_rate * 50))
        
        logger.info(f"\n📊 Keyword Score: {keyword_score}/50")
        logger.info(f"   Overall density: {keyword_stats['overall_density']:.1%}")
        logger.info(f"   Keywords matched: {len(keyword_stats['keywords'])}/{len(keywords)}")
        logger.info(f"   Excellent coverage: {excellent_count} keywords")
        logger.info(f"   Good coverage: {good_count} keywords")
        
        # Show top matches
        logger.info(f"\n   Top keyword matches:")
        sorted_keywords = sorted(
            keyword_stats['keywords'].items(),
            key=lambda x: x[1]['coverage'],
            reverse=True
        )[:10]
        
        for keyword, stats in sorted_keywords:
            status_icon = "✅" if stats['status'] == 'excellent' else "⚠️" if stats['status'] == 'good' else "❌"
            logger.info(f"   {status_icon} {keyword}: {stats['resume_count']}x in resume, {stats['jd_count']}x in JD ({stats['coverage']:.0f}% coverage)")
        
        # 3. Formatting/readability (30 points)
        format_score = 30
        
        # Check bullet point quality
        if structure_result['short_bullets'] > 0:
            format_score -= min(10, structure_result['short_bullets'] * 2)
        if structure_result['long_bullets'] > 0:
            format_score -= min(10, structure_result['long_bullets'] * 2)
        
        # Check keyword density
        if keyword_stats['overall_density'] < self.min_keyword_density:
            format_score -= 10
            logger.info(f"\n   ⚠️  Keyword density too low: {keyword_stats['overall_density']:.1%} < {self.min_keyword_density:.1%}")
        
        logger.info(f"\n📝 Formatting Score: {format_score}/30")
        
        # Calculate total
        total_score = structure_score + keyword_score + format_score
        
        logger.info(f"\n{'='*70}")
        logger.info(f"🎯 TOTAL ATS SCORE: {total_score}/100")
        logger.info(f"{'='*70}\n")
        
        return {
            'total_score': total_score,
            'breakdown': {
                'structure': structure_score,
                'keywords': keyword_score,
                'formatting': format_score
            },
            'structure_result': structure_result,
            'keyword_stats': keyword_stats,
            'recommendations': self._generate_recommendations(
                total_score,
                structure_result,
                keyword_stats
            )
        }
    
    def _generate_recommendations(
        self,
        score: int,
        structure_result: Dict,
        keyword_stats: Dict
    ) -> List[str]:
        """Generate improvement recommendations."""
        recommendations = []
        
        if score >= 90:
            recommendations.append("✅ Excellent! Your resume is highly optimized for ATS.")
        elif score >= 75:
            recommendations.append("⚠️  Good, but there's room for improvement.")
        else:
            recommendations.append("❌ Needs significant optimization for ATS.")
        
        # Structure recommendations
        for issue in structure_result['issues']:
            if 'Missing' in issue:
                recommendations.append(f"Add {issue.replace('Missing ', '')} to resume")
        
        # Keyword recommendations
        low_coverage_keywords = [
            k for k, v in keyword_stats['keywords'].items()
            if v['status'] == 'low' and v['jd_count'] >= 3
        ]
        
        if low_coverage_keywords:
            recommendations.append(
                f"Increase usage of: {', '.join(low_coverage_keywords[:5])}"
            )
        
        # Density recommendation
        if keyword_stats['overall_density'] < self.target_keyword_density:
            recommendations.append(
                f"Increase keyword density from {keyword_stats['overall_density']:.1%} to {self.target_keyword_density:.1%}"
            )
        
        return recommendations
    
    def compare_formatting(
        self,
        original_doc: Document,
        tailored_doc: Document
    ) -> Dict[str, any]:
        """
        Compare formatting between original and tailored resume.
        
        Returns:
            Dict with comparison results
        """
        logger.info(f"\n{'='*70}")
        logger.info(f"🔍 FORMATTING COMPARISON")
        logger.info(f"{'='*70}\n")
        
        differences = []
        
        # Compare paragraph counts
        if len(original_doc.paragraphs) != len(tailored_doc.paragraphs):
            differences.append(
                f"Paragraph count changed: {len(original_doc.paragraphs)} → {len(tailored_doc.paragraphs)}"
            )
        
        # Sample first 10 paragraphs for format comparison
        for i in range(min(10, len(original_doc.paragraphs), len(tailored_doc.paragraphs))):
            orig_para = original_doc.paragraphs[i]
            tail_para = tailored_doc.paragraphs[i]
            
            # Compare alignment
            if orig_para.alignment != tail_para.alignment:
                differences.append(f"Para {i}: Alignment changed")
            
            # Compare run formatting (if runs exist)
            if orig_para.runs and tail_para.runs:
                orig_run = orig_para.runs[0]
                tail_run = tail_para.runs[0]
                
                if orig_run.font.name != tail_run.font.name:
                    differences.append(f"Para {i}: Font changed from {orig_run.font.name} to {tail_run.font.name}")
                
                if orig_run.font.size != tail_run.font.size:
                    differences.append(f"Para {i}: Font size changed")
        
        formatting_preserved = len(differences) == 0
        
        if formatting_preserved:
            logger.info("✅ Formatting perfectly preserved!")
        else:
            logger.info(f"⚠️  {len(differences)} formatting differences found:")
            for diff in differences[:10]:
                logger.info(f"   • {diff}")
        
        logger.info(f"{'='*70}\n")
        
        return {
            'preserved': formatting_preserved,
            'differences': differences,
            'difference_count': len(differences)
        }
