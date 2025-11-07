"""
Resume management service for loading, tailoring, and saving resumes.
"""

from typing import Dict, Tuple, Optional, List
from pathlib import Path
from docx import Document
import copy
import re

from src.config.settings import Settings
from src.services.ai_service import AIService
from src.utils.logger import get_logger
from legacy_scripts.tailor_docx_resume import extract_resume_content, update_resume_sections

logger = get_logger(__name__)


class ResumeService:
    """Handle resume loading, tailoring, and saving operations."""
    
    def __init__(self, settings: Settings, ai_service: AIService):
        """
        Initialize resume service.
        
        Args:
            settings: Application settings
            ai_service: AI service for resume tailoring
        """
        self.settings = settings
        self.ai_service = ai_service
        self.resumes_dir = Path("resumes")
        self.original_resume_dir = self.resumes_dir / "original"
        self.tailored_resume_dir = self.resumes_dir / "tailored"
        
        # Ensure directories exist
        self.tailored_resume_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info("Resume service initialized")
    
    def load_original_resume(self) -> Tuple[str, Optional[Document]]:
        """
        Load the original resume from file.
        
        Returns:
            Tuple of (resume_text, document_object or None)
            - If DOCX: (text_content, Document object)
            - If TXT: (text_content, None)
        """
        resume_path = self.settings.user_info.files.original_resume_path
        try:
            if resume_path.endswith('.docx'):
                doc = Document(resume_path)
                text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                logger.info(f"Loaded DOCX resume: {resume_path}")
                return text, doc
            else:
                with open(resume_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                logger.info(f"Loaded TXT resume: {resume_path}")
                return text, None
        except FileNotFoundError:
            logger.error(f"❌ Resume file not found: {resume_path}")
            raise
    
    def tailor_resume(
        self,
        original_resume_text: str,
        job_description: str,
        job_title: str,
        company: str,
        original_doc: Optional[Document] = None
    ):
        """
        Use AI to tailor resume for specific job.
        
        Args:
            original_resume_text: Original resume content
            job_description: Full job description
            job_title: Job title
            company: Company name
            original_doc: Optional Document object for section-based tailoring
            
        Returns:
            Tailored resume (Document object if DOCX, string if TXT)
        """
        logger.info("="*70)
        logger.info("🤖 AI is tailoring your resume...")
        logger.info("="*70)
        
        # If we have a Document object, do section-by-section tailoring
        if original_doc:
            return self._tailor_docx_resume(original_doc, job_description, job_title, company)
        
        # Otherwise, do full text tailoring (legacy)
        return self._tailor_text_resume(original_resume_text, job_description, job_title, company)
    
    def _tailor_text_resume(
        self,
        resume_text: str,
        job_description: str,
        job_title: str,
        company: str
    ) -> str:
        """Tailor plain text resume."""
        prompt = f"""You are a professional resume writer. Rewrite the following resume to perfectly match this job posting.

JOB TITLE: {job_title}

JOB DESCRIPTION:
{job_description[:2000]}

ORIGINAL RESUME:
{resume_text}

Instructions:
1. Keep all factual information accurate (don't make up experience)
2. Emphasize skills and experiences that match the job requirements
3. Use keywords from the job description
4. Make it ATS-friendly
5. Keep it concise and professional
6. Format as plain text

Provide ONLY the tailored resume text, no additional commentary."""

        try:
            tailored_resume = self.ai_service.generate_text(prompt)
            logger.info(f"✅ Resume tailored successfully ({len(tailored_resume)} characters)")
            return tailored_resume
        except Exception as e:
            logger.error(f"❌ Error tailoring resume: {e}")
            logger.warning("   Using original resume as fallback")
            return resume_text
    
    def _tailor_docx_resume(
        self,
        doc: Document,
        job_description: str,
        job_title: str,
        company: str
    ) -> Document:
        """
        Tailor DOCX resume by updating specific sections while preserving formatting.
        
        Args:
            doc: Original Document object
            job_description: Full job description
            job_title: Job title
            company: Company name
            
        Returns:
            Document object with tailored content
        """
        logger.info("🔍 Extracting resume sections...")
        content = extract_resume_content(doc)
        
        # Tailor Summary section
        tailored_summary = ""
        if content['summary']['paras']:
            logger.info("\n📝 Tailoring SUMMARY section...")
            original_summary = '\n'.join([p['text'] for p in content['summary']['paras']])
            tailored_summary = self.ai_service.tailor_resume_summary(
                original_summary,
                job_description,
                job_title,
                company
            )
            logger.info("   ✅ Summary updated")
        
        # Tailor Skills section
        tailored_skills = ""
        if content['skills']['paras']:
            logger.info("\n🛠️  Tailoring TECHNICAL SKILLS section...")
            original_skills = '\n'.join([p['text'] for p in content['skills']['paras']])
            tailored_skills = self.ai_service.tailor_skills_section(
                original_skills,
                job_description,
                job_title,
                company
            )
            logger.info("   ✅ Skills updated")
        
        # Tailor Experience section (each job)
        tailored_jobs = []
        if content['experience_jobs']:
            logger.info(f"\n💼 Tailoring WORK EXPERIENCE ({len(content['experience_jobs'])} jobs)...")
            for idx, job in enumerate(content['experience_jobs'], 1):
                logger.info(f"\n   Job {idx}/{len(content['experience_jobs'])}: {job.get('company', 'Unknown')}")
                original_bullets = [p['text'] for p in job['paras']]
                
                job_info = {
                    'company': job.get('company', ''),
                    'title': job.get('title', ''),
                    'dates': job.get('dates', ''),
                    'bullets': original_bullets
                }
                
                tailored_bullets = self.ai_service.tailor_work_experience(
                    job_info,
                    job_description,
                    job_title,
                    company
                )
                tailored_jobs.append({'bullets': tailored_bullets})
                logger.info(f"      ✅ Updated {len(tailored_bullets)} bullet points")
        
        # Update the document
        logger.info(f"\n📝 Applying changes to document...")
        doc = update_resume_sections(doc, content, tailored_summary, tailored_skills, tailored_jobs)
        
        # Validate and enhance resume
        logger.info(f"\n🔍 Validating tailored resume...")
        doc = self._validate_and_enhance_resume(
            doc, content, tailored_summary, tailored_skills,
            tailored_jobs, job_description, job_title, company
        )
        
        logger.info(f"\n✅ Resume tailoring complete!")
        return doc
    
    def _validate_and_enhance_resume(
        self,
        doc: Document,
        content: Dict,
        tailored_summary: str,
        tailored_skills: str,
        tailored_jobs: List[Dict],
        job_description: str,
        job_title: str,
        company: str
    ) -> Document:
        """
        Validate tailored resume and enhance if sections are missing or incomplete.
        """
        issues_found = []
        
        # Check Summary section
        if not tailored_summary or len(tailored_summary.strip()) < 50:
            issues_found.append("Summary too short or missing")
            logger.warning("   ⚠️  Summary section incomplete")
        
        # Check Skills section
        if not tailored_skills or len(tailored_skills.strip()) < 20:
            issues_found.append("Skills section incomplete")
            logger.warning("   ⚠️  Skills section incomplete")
        else:
            # Check if all 4 categories are present
            required_categories = ['Programming & Frameworks:', 'Data & AI/ML Tools:', 
                                   'Cloud & DevOps:', 'Databases & APIs:']
            missing_categories = [cat for cat in required_categories if cat not in tailored_skills]
            if missing_categories:
                issues_found.append(f"Missing skill categories: {missing_categories}")
                logger.warning(f"   ⚠️  Missing categories: {', '.join(missing_categories)}")
        
        # Check Experience section
        if not tailored_jobs or len(tailored_jobs) == 0:
            issues_found.append("Experience section missing")
            logger.warning("   ⚠️  Experience section incomplete")
        else:
            for idx, job in enumerate(tailored_jobs, 1):
                if not job.get('bullets') or len(job.get('bullets', [])) == 0:
                    issues_found.append(f"Job {idx} has no bullets")
                    logger.warning(f"   ⚠️  Job {idx} missing bullet points")
        
        # If issues found, attempt to re-enhance
        if issues_found:
            logger.info(f"\n🔧 Enhancing incomplete sections...")
            
            # Re-enhance summary if needed
            if any('Summary' in issue for issue in issues_found):
                logger.info("   🔄 Re-generating summary...")
                original_summary = '\n'.join([p['text'] for p in content['summary']['paras']])
                tailored_summary = self.ai_service.tailor_resume_summary(
                    original_summary, job_description, job_title, company
                )
                content_updated = extract_resume_content(doc)
                doc = update_resume_sections(doc, content_updated, tailored_summary, tailored_skills, tailored_jobs)
            
            # Re-enhance skills if needed
            if any('Skills' in issue or 'categories' in issue for issue in issues_found):
                logger.info("   🔄 Re-generating skills...")
                original_skills = '\n'.join([p['text'] for p in content['skills']['paras']])
                tailored_skills = self.ai_service.tailor_skills_section(
                    original_skills, job_description, job_title, company
                )
                content_updated = extract_resume_content(doc)
                doc = update_resume_sections(doc, content_updated, tailored_summary, tailored_skills, tailored_jobs)
            
            logger.info("   ✅ Enhancement complete")
        else:
            logger.info("   ✅ All sections validated successfully")
        
        return doc
    
    def save_tailored_resume(
        self,
        company: str,
        job_title: str,
        tailored_content,
        original_doc: Optional[Document] = None
    ) -> str:
        """
        Save tailored resume to file in organized directory structure.
        Structure: tailored/{company}_{job_title}/{original_resume_name}.docx
        
        Args:
            company: Company name
            job_title: Job title
            tailored_content: Tailored resume (Document object or text string)
            original_doc: Original Document object (for text-based tailoring)
            
        Returns:
            Path to saved resume file
        """
        # Create safe folder name from company and job title
        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')
        folder_name = f"{safe_company}_{safe_title}"
        
        # Create company/job folder inside tailored directory
        job_folder = self.tailored_resume_dir / folder_name
        job_folder.mkdir(parents=True, exist_ok=True)
        
        # Get original resume filename (without path)
        original_resume_path = Path(self.settings.user_info.files.original_resume_path)
        original_filename = original_resume_path.name
        
        # Check if it's a Document object
        if hasattr(tailored_content, 'save') and hasattr(tailored_content, 'paragraphs'):
            # Save as DOCX using original filename
            filepath = job_folder / original_filename
            tailored_content.save(str(filepath))
        else:
            # Text content
            if original_doc:
                # Save as DOCX (legacy simple approach)
                filepath = job_folder / original_filename
                
                new_doc = copy.deepcopy(original_doc)
                resume_lines = tailored_content.split('\n')
                
                # Clear and rebuild
                while len(new_doc.paragraphs) > 0:
                    p = new_doc.paragraphs[0]
                    p._element.getparent().remove(p._element)
                
                for line in resume_lines:
                    if line.strip():
                        new_doc.add_paragraph(line)
                
                new_doc.save(str(filepath))
            else:
                # Save as TXT using original filename
                txt_filename = original_filename.replace('.docx', '.txt')
                filepath = job_folder / txt_filename
                
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(tailored_content)
        
        logger.info(f"💾 Saved tailored resume: {filepath}")
        return str(filepath)
    
    def generate_resume_from_template(
        self,
        job_description: str,
        job_title: str,
        company: str
    ) -> str:
        """
        Generate a complete resume from template by analyzing job description.
        
        Args:
            job_description: Job description text
            job_title: Job title
            company: Company name
            
        Returns:
            Path to generated resume file
        """
        logger.info(f"\n🎯 Generating resume from template for {job_title} at {company}")
        
        # Load template
        template_path = self.original_resume_dir / "Resume.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        doc = Document(str(template_path))
        
        # Get personal info from config (build from user_info if resume_personal_info not set)
        if self.settings.user_info.resume_personal_info:
            personal_info = self.settings.user_info.resume_personal_info.model_dump()
        else:
            # Build from existing user_info
            personal_info = {
                'full_name': self.settings.user_info.personal_info.name,
                'location': f"{self.settings.user_info.personal_info.city}, {self.settings.user_info.personal_info.state_abbr}",
                'phone': self.settings.user_info.personal_info.phone,
                'email': self.settings.user_info.personal_info.email,
                'linkedin': self.settings.user_info.links.linkedin.replace('https://www.linkedin.com/in/', ''),
                'github': self.settings.user_info.links.github.replace('https://github.com/', ''),
                'title': self.settings.user_info.background.current_title,
                'years_experience': self.settings.user_info.background.years_of_experience,
                'core_stack': 'Python, React, AWS, MLOps',
                'value_statement': self.settings.user_info.background.elevator_pitch[:100]
            }
        
        # Generate each section using AI
        logger.info("📝 Generating header...")
        header = self._generate_header(personal_info, job_description, job_title, company)
        logger.info(f"HEADER AI RESPONSE: {header}")

        logger.info("🔧 Generating skills...")
        skills = self._generate_skills(job_description, job_title, company)
        logger.info(f"SKILLS AI RESPONSE: {skills}")

        logger.info("💼 Generating work experience...")
        # Use real work experience from config and generate bullets for each
        if self.settings.user_info.work_experience:
            experience = self._generate_experience_with_config(
                job_description, job_title, company, self.settings.user_info.work_experience
            )
        else:
            experience = self._generate_experience(job_description, job_title, company)
        logger.info(f"WORK EXPERIENCE AI RESPONSE: {experience}")

        logger.info("🎓 Using education from config...")
        # Use real education from config
        if self.settings.user_info.education:
            education = {
                'degree': self.settings.user_info.education.degree,
                'graduated': self.settings.user_info.education.graduated,
                'gpa': self.settings.user_info.education.gpa or '',
                'university': self.settings.user_info.education.university,
                'location': self.settings.user_info.education.location,
                'coursework': self.settings.user_info.education.coursework or 'Data Structures, Algorithms, Database Systems, Operating Systems, Computer Networks, Machine Learning'
            }
        else:
            education = self._generate_education()
        logger.info(f"EDUCATION AI RESPONSE: {education}")
        
        # Replace template placeholders
        self._fill_template(doc, header, skills, experience, education)
        
        # Save generated resume
        safe_company = "".join(c for c in (company or 'Company') if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')
        safe_title = "".join(c for c in (job_title or 'Position') if c.isalnum() or c in (' ', '_')).strip().replace(' ', '_')
        
        output_dir = self.tailored_resume_dir / f"{safe_company}_{safe_title}"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename: FirstName_LastName_CompanyName_Resume.docx
        first_name = personal_info['full_name'].split()[0] if personal_info['full_name'] else 'Resume'
        last_name = personal_info['full_name'].split()[-1] if len(personal_info['full_name'].split()) > 1 else ''
        filename = f"{first_name}_{last_name}_{safe_company}_Resume.docx" if last_name else f"{first_name}_{safe_company}_Resume.docx"
        
        output_file = output_dir / filename
        doc.save(str(output_file))
        
        logger.info(f"✅ Generated resume: {output_file}")
        return str(output_file)
    
    def _generate_header(self, personal_info: dict, job_description: str, job_title: str, company: str) -> dict:
        """Generate professional header matching job requirements."""
        
        # Calculate years of experience from work history with range of 8-10
        years_experience = 10  # Default fallback
        if self.settings.user_info.work_experience:
            # Calculate from first job start date to current
            try:
                first_job = self.settings.user_info.work_experience[-1]  # Oldest job (last in list)
                # Extract start year from dates like "Feb 2015 – Jan 2019"
                start_date = first_job.dates.split('–')[0].strip()
                start_year = int(start_date.split()[-1])
                current_year = 2025
                calculated_years = current_year - start_year
                years_experience = min(10, max(8, calculated_years))  # Between 8 and 10
            except:
                years_experience = 10  # Fallback to 10 if parsing fails
        
        # Analyze job to determine professional title and core stack
        prompt = f"""Based on this job description for {job_title} at {company}, generate a professional resume header.

Job Description:
{job_description[:1500]}

Generate EXACTLY 4 lines (no more, no less), each with ONLY the requested content:

Line 1: Professional title (e.g., "Senior Backend Engineer" or "Full Stack Developer" or "ML Engineer")
Line 2: Years of experience as a number with + (e.g., "10+" or "8+" or "9+") - MUST be between 8 and 10
Line 3: Core tech stack - list 3-5 key technologies from job description, comma-separated (e.g., "Python, React, AWS, PostgreSQL")
Line 4: Value statement - one impressive achievement in 10-15 words (e.g., "Delivering scalable cloud solutions with 99.9% uptime")

CRITICAL RULES:
- Output EXACTLY 4 lines
- Line 2 MUST be between 8 and 10 with + symbol (e.g., "8+", "9+", or "10+")
- NO additional text, commentary, explanations, or prefixes
- NO markdown formatting, NO asterisks, NO bullets
- NO phrases like "Here is", "Response:", or any headers
- Just the 4 values, one per line

Example output:
Senior Software Engineer
10+
Ruby on Rails, React, AWS, PostgreSQL
Building high-performance systems serving 10M+ users daily
"""

        response = self.ai_service.generate_completion(prompt).strip()
        
        # Clean AI response aggressively
        response = response.replace('**', '').replace('*', '')
        # Remove common AI prefixes from entire response
        for prefix in ['here is the generated resume header:', 'here is:', 'here are:', 'response:', 'output:']:
            if response.lower().startswith(prefix):
                response = response[len(prefix):].strip()
        
        lines = [line.strip() for line in response.split('\n') if line.strip() and not line.strip().startswith(('line ', 'Line '))]
        
        # Extract values with validation
        title = (lines[0] if len(lines) > 0 else job_title).upper()
        
        # Extract and validate years - ensure between 8 and 10
        years_raw = lines[1] if len(lines) > 1 else f"{years_experience}+"
        # Parse the number from AI response (e.g., "3+" -> 3)
        try:
            years_num = int(years_raw.rstrip('+').strip())
            # Clamp between 8 and 10
            years_num = min(10, max(8, years_num))
            years = f"{years_num}+"
        except:
            years = f"{years_experience}+"
        
        core_stack = lines[2] if len(lines) > 2 else "Full Stack Development"
        value_statement = lines[3] if len(lines) > 3 else "Delivering robust, scalable solutions for business impact"
        
        return {
            'full_name': personal_info['full_name'],
            'location': personal_info['location'],
            'phone': personal_info['phone'],
            'email': personal_info['email'],
            'linkedin': personal_info['linkedin'],
            'github': personal_info['github'],
            'title': title,
            'years': years,
            'core_stack': core_stack,
            'value_statement': value_statement
        }
    
    def _generate_skills(self, job_description: str, job_title: str, company: str) -> dict:
        """Generate skills section based on job requirements."""
        
        prompt = f"""
Analyze this job description for {job_title} at {company} and generate a comprehensive Technical Skills section for a resume.

Job Description:
{job_description[:2000]}

Generate skills in these exact categories (one line per category):
- Languages: [programming languages]
- AI/ML: [ML frameworks, tools]
- Frontend: [frontend frameworks and tools]
- Backend: [backend frameworks, APIs, architectures]
- Databases: [databases used]
- Cloud & DevOps: [cloud platforms, DevOps tools]
- Other: [other relevant technologies]

CRITICAL RULES:
1. For EVERY category, provide a concrete list of technologies - NEVER use "None", "N/A", "not mentioned", or vague terms
2. If job description doesn't mention specific technologies for a category, use industry-standard defaults:
   - Languages: Python, JavaScript, Java, TypeScript
   - AI/ML: TensorFlow, PyTorch, scikit-learn
   - Frontend: React, Next.js, Angular, Vue.js
   - Backend: Node.js, Django, Spring Boot, Express
   - Databases: PostgreSQL, MySQL, MongoDB, Redis
   - Cloud & DevOps: AWS, Docker, Kubernetes, CI/CD
   - Other: Git, REST APIs, GraphQL, Microservices
3. Output ONLY the 7 lines in format "Category: tech1, tech2, tech3"
4. NO explanations, NO parentheses, NO commentary
5. Be specific - use actual technology names, not generic terms

Example output:
Languages: Python, JavaScript, Java
AI/ML: TensorFlow, PyTorch, scikit-learn
Frontend: React, Next.js, TypeScript
Backend: Node.js, Django, FastAPI
Databases: PostgreSQL, MongoDB, Redis
Cloud & DevOps: AWS, Docker, Kubernetes
Other: Git, REST APIs, GraphQL
"""

        response = self.ai_service.generate_completion(prompt).strip()
        
        # Remove empty lines and clean
        lines = [line.strip() for line in response.split('\n') if line.strip() and ':' in line]

        skills_dict = {}
        for line in lines:
            if ':' in line:
                category, items = line.split(':', 1)
                cat = category.strip()
                val = items.strip()
                
                # Aggressive validation: if value is None/N/A/empty, use defaults
                # Check both the raw value and lowercase version
                if not val or len(val) < 3 or val.lower() in ['none', 'n/a', 'not mentioned', 'not specified', 'not applicable']:
                    # Provide sensible defaults based on category
                    if 'language' in cat.lower():
                        val = 'Python, JavaScript, Java, TypeScript'
                    elif 'ai' in cat.lower() or 'ml' in cat.lower():
                        val = 'TensorFlow, PyTorch, scikit-learn'
                    elif 'frontend' in cat.lower():
                        val = 'React, Next.js, Angular, Vue.js'
                    elif 'backend' in cat.lower():
                        val = 'Node.js, Django, Spring Boot, Express'
                    elif 'database' in cat.lower():
                        val = 'PostgreSQL, MySQL, MongoDB, Redis'
                    elif 'cloud' in cat.lower() or 'devops' in cat.lower():
                        val = 'AWS, Docker, Kubernetes, CI/CD'
                    else:
                        val = 'Git, REST APIs, GraphQL, Microservices'
                
                skills_dict[cat] = val

        return skills_dict
    
    def _generate_experience(self, job_description: str, job_title: str, company: str) -> list:
        """Generate work experience bullets for past jobs."""
        
        # Get work history from config
        prompt = f"""Generate impressive work experience bullets for a resume targeting {job_title} at {company}.

Job Description (what they want):
{job_description[:2000]}

Generate 4 past job experiences, each with:
- Job Title | Company Name | Location | Dates (e.g., "Dec 2024 — Current")
- 4-5 bullet points following this formula:
  [Action Verb] [what you did] using [technology from job description], resulting in [quantified impact with metrics]

Rules:
1. Use technologies and skills from the job description
2. Every bullet MUST have quantified metrics (%, $, time saved, users, etc.)
3. Start with strong action verbs: Architected, Engineered, Optimized, Built, Deployed, Reduced, Increased
4. Show progression: most recent job is most impressive
5. Each bullet should be 1-2 lines max
6. Make dates realistic (work backwards from current date)

Format each job as:
JOB_TITLE | COMPANY | LOCATION | DATES
• Bullet point 1
• Bullet point 2
• Bullet point 3
• Bullet point 4

Generate all 4 jobs now:"""

        response = self.ai_service.generate_completion(prompt, temperature=0.8).strip()
        
        # Parse the response into structured format
        jobs = []
        current_job = None
        
        for line in response.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            # Check if it's a job header (contains |)
            if '|' in line and not line.startswith('•'):
                if current_job:
                    jobs.append(current_job)
                
                parts = [p.strip() for p in line.split('|')]
                current_job = {
                    'title': parts[0] if len(parts) > 0 else '',
                    'company': parts[1] if len(parts) > 1 else '',
                    'location': parts[2] if len(parts) > 2 else '',
                    'dates': parts[3] if len(parts) > 3 else '',
                    'bullets': []
                }
            elif line.startswith('•') and current_job:
                bullet = line.lstrip('•').strip()
                current_job['bullets'].append(bullet)
        
        if current_job:
            jobs.append(current_job)
        
        return jobs[:4]  # Return max 4 jobs
    
    def _generate_experience_with_config(self, job_description: str, job_title: str, company: str, work_history: list) -> list:
        """Generate work experience bullets using real work history from config."""
        
        jobs = []
        for work in work_history:
            # Extract year range from dates (e.g., "Feb 2015 – Jan 2019" -> 2015-2019)
            try:
                dates_parts = work.dates.replace('–', '-').replace('—', '-').split('-')
                start_year = int(dates_parts[0].strip().split()[-1])
                if len(dates_parts) > 1 and 'Current' not in dates_parts[1]:
                    end_year = int(dates_parts[1].strip().split()[-1])
                else:
                    end_year = 2025  # Current
            except:
                start_year = 2020
                end_year = 2025
            
            # Generate bullets for this specific job
            prompt = f"""Generate 4-5 impressive work experience bullets for this role, tailored to the target job.

TARGET JOB YOU'RE APPLYING FOR:
Position: {job_title} at {company}
Description: {job_description[:2000]}

YOUR PAST ROLE (what you accomplished):
Title: {work.title}
Company: {work.company}
Location: {work.location}
Dates: {work.dates} (Years: {start_year}-{end_year})

CRITICAL RULES:
1. You are generating bullets for YOUR PAST work at {work.company}, NOT for the company you're applying to ({company})
2. NEVER mention {company} as a tool, platform, or technology - it's the company you're APPLYING TO, not a tool you used
3. Use relevant technologies from the target job description, BUT only use technologies that existed during {start_year}-{end_year}
4. TECHNOLOGY TIMELINE - Only use technologies released BEFORE or DURING this time period:
   - React (2013+), Angular (2010+), Vue.js (2014+), Next.js (2016+), Node.js (2009+)
   - Python (always available), Java (always available), JavaScript (always available)
   - Docker (2013+), Kubernetes (2014+), AWS (2006+), Azure (2010+), GCP (2008+)
   - MongoDB (2009+), PostgreSQL (always available), Redis (2009+), MySQL (always available)
   - TensorFlow (2015+), PyTorch (2016+), scikit-learn (2007+)
   - TypeScript (2012+), Go (2009+), Rust (2010+)
5. DO NOT use technologies that didn't exist during {start_year}-{end_year} (e.g., don't use Next.js for 2010-2014)
6. Every bullet MUST have quantified metrics (%, $, time saved, users, etc.)
7. Start with strong action verbs: Architected, Engineered, Optimized, Built, Deployed, Reduced, Increased
8. DO NOT use "Led", "Lead", or "Leading" as action verbs - focus on technical implementation verbs
9. Each bullet should be 1-2 lines max
10. Show how your past work at {work.company} makes you qualified for {job_title} at {company}

GRAMMAR AND FORMATTING RULES (CRITICAL):
11. Use correct past tense verbs: "Architected" NOT "Architectured", "Built" NOT "Builded"
12. For verb consistency: "views, processing" NOT "views, processed" (use gerund after comma)
13. Every bullet MUST end with a period (.)
14. Perfect spelling and grammar - proofread each bullet
15. No double periods (..)
16. Proper spacing around punctuation
17. No commentary, explanations, or meta-text like "Here is" or "Corrected:"
18. Output ONLY the bullet text, starting directly with the action verb

Example format for {start_year}-{end_year}:
• Built scalable microservices using React and AWS, reducing latency by 40% and serving 2M+ users.
• Architected data pipeline with Python and PostgreSQL, processing 50TB daily with 99.9% reliability.

Output ONLY the bullets (4-5 bullets), one per line, starting with •
Each bullet must be grammatically perfect with proper punctuation.
NO additional commentary, explanations, prefixes, or meta-text."""

            response = self.ai_service.generate_completion(prompt, temperature=0.8).strip()
            
            # Parse bullets - basic cleanup only
            bullets = []
            for line in response.split('\n'):
                line = line.strip()
                if line.startswith('•'):
                    bullet = line.lstrip('•').strip()
                    
                    # Basic cleanup - ensure period at end
                    if bullet and not bullet.endswith(('.', '!', '?')):
                        bullet += '.'
                    
                    # Filter out bullets that:
                    # 1. Mention target company as a tool
                    # 2. Start with "Led" or "Lead"
                    if (company.lower() not in bullet.lower() or f"at {company}" in bullet.lower()) and \
                       not bullet.lower().startswith(('led ', 'lead ')):
                        bullets.append(bullet)
            
            if bullets:
                jobs.append({
                    'title': work.title,
                    'company': work.company,
                    'location': work.location,
                    'dates': work.dates,
                    'bullets': bullets[:5]  # Max 5 bullets
                })
        
        return jobs
    
    def _generate_education(self) -> dict:
        """Generate education section."""
        # For now, use a template - can be enhanced later
        return {
            'degree': 'Bachelor of Science in Computer Science',
            'graduated': 'May 2015',
            'gpa': '3.8/4.0',
            'university': 'University of Technology',
            'location': 'Boston, MA',
            'coursework': 'Data Structures, Algorithms, Database Systems, Operating Systems, Computer Networks, Machine Learning'
        }
    
    def _fill_template(self, doc: Document, header: dict, skills: dict, experience: list, education: dict):
        """Fill the template with generated content while preserving formatting."""
        
        for paragraph in doc.paragraphs:
            text = paragraph.text
            # Replace header placeholders while preserving formatting
            if '[Full Name]' in text:
                self._replace_text_in_paragraph(paragraph, '[Full Name]', header['full_name'])
            elif '[Location]' in text and 'xxx' in text:
                self._replace_text_in_paragraph(paragraph, text, f"{header['location']} | {header['phone']} | {header['email']}")
            elif 'LinkedIn:' in text and 'xxxxx' in text:
                self._replace_text_in_paragraph(paragraph, text, f"LinkedIn: {header['linkedin']} | GitHub: {header['github']}")
            elif 'PROFESSIONAL TITLE' in text:
                # Compose the professional title block in all uppercase and clean
                title_block = f"{header['title']} | {header['years']} | {header['core_stack']} | {header['value_statement']}"
                for prefix in ["HERE ARE THE REQUESTED ELEMENTS:", "HERE IS:", "HERE ARE:", "RESPONSE:", "|", ":"]:
                    if title_block.strip().upper().startswith(prefix):
                        title_block = title_block.strip()[len(prefix):].strip()
                title_block = title_block.upper()
                self._replace_text_in_paragraph(paragraph, text, title_block)
            
            # Replace skills placeholders using actual AI response keys, stripping markdown and fallback to correct keys if needed
            def clean_skill_value(val):
                if not val:
                    return ''
                val = val.replace('**', '').replace('*', '').strip()
                # Check if it's a placeholder value that should be ignored
                if val.lower() in ['none', 'n/a', 'not mentioned', 'not specified', 'not applicable']:
                    return ''
                # Remove parentheses and explanations
                val = re.sub(r'\(.*?\)', '', val)
                val = val.replace('e.g.', '').replace('E.g.', '').replace('for example', '').replace('no mention of', '').replace('implied to be', '').replace('implied', '').replace('but not specified', '').replace('not specified', '').replace('N/A', '').replace('None', '').replace(':', '').strip()
                # Remove extra spaces and commas
                val = re.sub(r',\s*,', ',', val)
                val = re.sub(r'\s{2,}', ' ', val)
                val = val.strip(', ').strip()
                # Final check - if after cleaning it's too short or empty, return empty
                if len(val) < 3:
                    return ''
                return val

            def get_skill_value(skills, key, fallback=None):
                val = skills.get(key)
                if not val:
                    val = skills.get(f"**{key}")
                val = clean_skill_value(val)
                # If cleaned value is empty or None-like, use fallback
                if not val or len(val) < 3:
                    val = fallback if fallback else ''
                return val

            # Widely used tech for each category
            default_skills = {
                'Languages': 'Python, JavaScript, TypeScript, Java, Kotlin',
                'AI/ML': 'TensorFlow, PyTorch, scikit-learn, OpenCV',
                'Frontend': 'React, Next.js, Angular, Vue.js',
                'Backend': 'Node.js, Flask, Django, Spring Boot',
                'Databases': 'PostgreSQL, MongoDB, MySQL, Redis',
                'Cloud & DevOps': 'AWS, Docker, Kubernetes, Azure, GCP',
                'Other': 'CI/CD, Git, REST APIs, Agile, Functional Programming'
            }

            if text.startswith('Languages:') and '[List' in text:
                value = get_skill_value(skills, 'Languages', default_skills['Languages'])
                self._replace_text_in_paragraph(paragraph, text, f"Languages: {value}")
            elif text.startswith('AI/ML') and '[List' in text:
                value = get_skill_value(skills, 'AI/ML', default_skills['AI/ML'])
                self._replace_text_in_paragraph(paragraph, text, f"AI/ML: {value}")
            elif text.startswith('Frontend:') and '[List' in text:
                value = get_skill_value(skills, 'Frontend', default_skills['Frontend'])
                self._replace_text_in_paragraph(paragraph, text, f"Frontend: {value}")
            elif text.startswith('Backend:') and '[List' in text:
                value = get_skill_value(skills, 'Backend', default_skills['Backend'])
                self._replace_text_in_paragraph(paragraph, text, f"Backend: {value}")
            elif text.startswith('Databases:') and '[List' in text:
                value = get_skill_value(skills, 'Databases', default_skills['Databases'])
                self._replace_text_in_paragraph(paragraph, text, f"Databases: {value}")
            elif text.startswith('Cloud & DevOps:') and '[List' in text:
                value = get_skill_value(skills, 'Cloud & DevOps', default_skills['Cloud & DevOps'])
                self._replace_text_in_paragraph(paragraph, text, f"Cloud & DevOps: {value}")
            elif text.startswith('Other:') and '[List' in text:
                value = get_skill_value(skills, 'Other', default_skills['Other'])
                self._replace_text_in_paragraph(paragraph, text, f"Other: {value}")
        
        # Find and delete all template placeholders between RELEVANT WORK EXPERIENCE and EDUCATION
        work_header_idx = None
        education_idx = None
        for i, paragraph in enumerate(doc.paragraphs):
            if 'RELEVANT WORK EXPERIENCE' in paragraph.text:
                work_header_idx = i
            if 'EDUCATION' in paragraph.text and education_idx is None:
                education_idx = i
        
        # Delete all paragraphs between RELEVANT WORK EXPERIENCE and EDUCATION (template placeholders)
        if work_header_idx is not None and education_idx is not None:
            # Get all paragraph elements between these indices
            paras_to_delete = []
            for i in range(work_header_idx + 1, education_idx):
                if i < len(doc.paragraphs):
                    paras_to_delete.append(doc.paragraphs[i])
            
            # Delete them
            for para in paras_to_delete:
                p_element = para._element
                p_element.getparent().remove(p_element)
        
        # Re-find indices after deletion
        work_header_idx = None
        education_idx = None
        for i, paragraph in enumerate(doc.paragraphs):
            if 'RELEVANT WORK EXPERIENCE' in paragraph.text:
                work_header_idx = i
            if 'EDUCATION' in paragraph.text and education_idx is None:
                education_idx = i
        
        # Insert each job experience block in correct order (most recent first)
        # Config has jobs in correct order: [0]=Ntiva (newest), [1]=Insight, [2]=Codoxo, [3]=Plego (oldest)
        # We insert before EDUCATION section
        if education_idx is not None:
            education_para = doc.paragraphs[education_idx]
            education_element = education_para._element
            parent = education_element.getparent()
            
            # Insert each job header first (bolded), then bullets below, for each job in config order
            for job in experience:
                # Insert job header (bold)
                new_p = parent.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', nsmap=education_element.nsmap)
                parent.insert(parent.index(education_element), new_p)
                from docx.text.paragraph import Paragraph
                header_para = Paragraph(new_p, parent)
                run = header_para.add_run(f"{job['title']} | {job['company']} | {job['location']} | {job['dates']}")
                run.bold = True

                # Insert all bullets (in order)
                for bullet in job['bullets']:
                    clean_bullet = bullet.replace('**', '').replace('*', '').strip()
                    new_p = parent.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', nsmap=education_element.nsmap)
                    parent.insert(parent.index(education_element), new_p)
                    bullet_para = Paragraph(new_p, parent)
                    bullet_para.add_run(f"• {clean_bullet}")

        # Bold section headers and key fields
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().upper() in [
                'TECHNICAL SKILLS', 'RELEVANT WORK EXPERIENCE', 'EDUCATION', 'CERTIFICATIONS'
            ]:
                for run in paragraph.runs:
                    run.bold = True
            # Bold key fields in education and certifications
            if paragraph.text.strip().startswith('Bachelor') or paragraph.text.strip().startswith('Master'):
                for run in paragraph.runs:
                    run.bold = True
            if 'Certified' in paragraph.text or 'Certification' in paragraph.text:
                for run in paragraph.runs:
                    run.bold = True
            if paragraph.text.strip().startswith('Relevant Coursework:'):
                for run in paragraph.runs:
                    run.bold = True
            
            # Clear template instructions and replace education placeholders
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if '[Focus on impact' in text or '[Include collaboration' in text or '[Always quantify' in text:
                    paragraph.clear()
                elif '[Each bullet' in text or '[Remember' in text or '[Template' in text:
                    paragraph.clear()
                elif '[Degree Name]' in text:
                    self._replace_text_in_paragraph(paragraph, '[Degree Name]', education['degree'])
                elif 'Graduated:' in text and '[Month Year]' in text:
                    if education.get('gpa'):
                        self._replace_text_in_paragraph(paragraph, text, f"Graduated: {education['graduated']} GPA: {education['gpa']}")
                    else:
                        self._replace_text_in_paragraph(paragraph, text, f"Graduated: {education['graduated']}")
                elif '[University Name]' in text:
                    self._replace_text_in_paragraph(paragraph, '[University Name]', education['university'])
                elif text == '[City, State]':
                    self._replace_text_in_paragraph(paragraph, '[City, State]', education['location'])
                elif 'Relevant Coursework:' in text and 'Data Structures' in text:
                    self._replace_text_in_paragraph(paragraph, text, f"Relevant Coursework: {education['coursework']}")
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """Replace text in paragraph while preserving formatting."""
        if old_text in paragraph.text:
            # Store original formatting
            if paragraph.runs:
                original_font = paragraph.runs[0].font
                paragraph.clear()
                run = paragraph.add_run(new_text)
                # Apply original formatting
                run.font.name = original_font.name
                run.font.size = original_font.size
                run.font.bold = original_font.bold
                run.font.italic = original_font.italic
            else:
                paragraph.text = new_text
